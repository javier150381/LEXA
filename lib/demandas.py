import os
import shutil
import glob
import json
import re
import unicodedata
import sqlite3
import sys
from pathlib import Path
from tkinter import messagebox, filedialog
from dotenv import load_dotenv

from .deepseek_key import get_deepseek_api_key

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from .pdf_utils import read_pdf_text

try:
    from pdf2image import convert_from_path
    import pytesseract
except Exception:  # noqa: BLE001 - optional OCR dependencies may not be present
    convert_from_path = None
    pytesseract = None
from langchain_community.vectorstores import FAISS
from email.header import decode_header
from langchain.schema import Document
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain.prompts import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.chains import ConversationalRetrievalChain
from langchain_core.output_parsers import StrOutputParser
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.retrievers import EnsembleRetriever

from langchain_openai import ChatOpenAI

from . import tokens
from .exact_index import build_exact_index, search_article
from lib import database as db
from src.schema_utils import (
    generate_schema_from_pdf,
    hash_for_pdf,
    index_hash,
    update_schema_index,
    cache_form_data,
    SCHEMAS_DIR,
)


class TrackingChatOpenAI(ChatOpenAI):
    """ChatOpenAI subclass that records token usage."""

    def _generate(self, messages, stop=None, run_manager=None, **kwargs):
        result = super()._generate(messages, stop=stop, run_manager=run_manager, **kwargs)
        try:
            usage = result.llm_output.get("token_usage", {})
            tokens_used = usage.get("total_tokens")
            prompt_tokens = usage.get("prompt_tokens")
            completion_tokens = usage.get("completion_tokens")
            if tokens_used:
                tokens.add_tokens(
                    int(tokens_used),
                    tokens_in=int(prompt_tokens) if prompt_tokens is not None else None,
                    tokens_out=int(completion_tokens) if completion_tokens is not None else None,
                )
        except Exception:
            pass
        return result

# Determine application base directory
def _get_base_dir():
    """Return a writable base directory for application data."""
    if getattr(sys, "frozen", False):
        user_dir = os.getenv("LOCALAPPDATA", os.path.expanduser("~"))
        return os.path.join(user_dir, "AbogadoVirtual")

    base = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    if not os.access(base, os.W_OK):
        user_dir = os.getenv("LOCALAPPDATA", os.path.expanduser("~"))
        return os.path.join(user_dir, "AbogadoVirtual")
    return base

BASE_DIR = _get_base_dir()

# Load environment variables from .env at repository root
load_dotenv()

DATA_DIR       = os.path.join(BASE_DIR, "data")
JURIS_DIR      = os.path.join(DATA_DIR, "jurisprudencia")
CASOS_DIR_ROOT = os.path.join(DATA_DIR, "casos")
DEMANDAS_DIR   = os.path.join(DATA_DIR, "demandas_ejemplo")
AREAS_DIR_ROOT = os.path.join(DATA_DIR, "areas")
LEGAL_CORPUS_DIR = os.path.join(DATA_DIR, "legal_corpus")

VECTOR_DB_JURIS    = os.path.join(BASE_DIR, "vector_db_juris")
VECTOR_DB_CASOS    = os.path.join(BASE_DIR, "vector_db_casos")
VECTOR_DB_DEMANDAS = os.path.join(BASE_DIR, "vector_db_demandas")
VECTOR_DB_LEGAL    = os.path.join(BASE_DIR, "vector_db_legal")

CONFIG_PATH = os.path.join(BASE_DIR, "config.json")

DB_PATH = os.path.join(DATA_DIR, "abogado.db")

ALERTAS_DB = os.path.join(DATA_DIR, "alertas.db")


def _resolve_path(path_value, default):
    """Return absolute path resolved against BASE_DIR when blank or relative."""
    if not path_value:
        return default
    if not os.path.isabs(path_value):
        return os.path.abspath(os.path.join(BASE_DIR, path_value))
    return path_value


for carpeta in (
    JURIS_DIR,
    CASOS_DIR_ROOT,
    DEMANDAS_DIR,
    AREAS_DIR_ROOT,
    LEGAL_CORPUS_DIR,
    VECTOR_DB_JURIS,
    VECTOR_DB_CASOS,
    VECTOR_DB_DEMANDAS,
    VECTOR_DB_LEGAL,
):
    os.makedirs(carpeta, exist_ok=True)


_llm = None
_llm_provider = None

def get_llm():
    """Return a cached ChatOpenAI instance if there is available credit."""
    tokens.init_db()
    if tokens.get_credit() <= 0:
        raise ValueError("Crédito agotado")

    global _llm, _llm_provider
    provider = getattr(default_context, "llm_provider", "deepseek")
    if _llm is None or _llm_provider != provider:
        _llm_provider = provider
        if provider == "openai":
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("❌ No se encontró OPENAI_API_KEY en .env")
            _llm = TrackingChatOpenAI(
                model="gpt-4o-mini",
                api_key=api_key,
                temperature=0.0,
            )
        else:
            api_key = get_deepseek_api_key()
            if not api_key:
                raise ValueError("❌ No se encontró DEEPSEEK_API_KEY en .env")
            _llm = TrackingChatOpenAI(
                model="deepseek-chat",
                base_url="https://api.deepseek.com/v1",
                api_key=api_key,
                temperature=0.0,
            )
    return _llm


def get_qa_prompt():
    system_msg = SystemMessagePromptTemplate.from_template(
        "Eres un abogado virtual de Ecuador. Usa los documentos cargados para responder en español."
    )
    human_msg = HumanMessagePromptTemplate.from_template(
        "Contexto legal:\n{context}\nPregunta:\n{question}"
    )
    return ChatPromptTemplate.from_messages([system_msg, human_msg])


prompt = get_qa_prompt()


def get_summary_prompt():
    system_msg = SystemMessagePromptTemplate.from_template(
        "Eres un abogado virtual de Ecuador. Usa los documentos cargados para responder en español."
    )
    human_msg = HumanMessagePromptTemplate.from_template(
        "Contexto legal:\n{context}\nInstrucciones:\n{question}"
    )
    return ChatPromptTemplate.from_messages([system_msg, human_msg])


summary_prompt = get_summary_prompt()


# Template basado en el artículo 142 del COGEP
COGEP_TEMPLATE = """
PRIMERO. - DESIGNACIÓN DEL JUZGADOR:
[DESIGNACION_JUZGADOR]
País: [PAIS]
Zona: [ZONA]

SEGUNDO. - DATOS DEL ACTOR:
Nombres y apellidos: [ACTOR_NOMBRES_APELLIDOS]
Cédula: [ACTOR_CEDULA]
Pasaporte: [ACTOR_PASAPORTE]
Estado civil: [ACTOR_ESTADO_CIVIL]
Edad: [ACTOR_EDAD]
Profesión: [ACTOR_PROFESION]
Provincia: [ACTOR_PROVINCIA]
Cantón: [ACTOR_CANTON]
Calle primaria: [ACTOR_CALLE_PRIMARIA]
Calle secundaria: [ACTOR_CALLE_SECUNDARIA]
Número de casa: [ACTOR_NUMERO_CASA]
Dirección electrónica: [ACTOR_DIR_ELECTRONICA]

TERCERO. - DATOS DEL DEFENSOR:
Nombre: [DEFENSOR_NOMBRE]
Casillero judicial: [CASILLERO_JUDICIAL]
Comparece como procurador: [REPRESENTA_COMO_PROCURADOR]
Datos del representado: [DATOS_REPRESENTADO]

CUARTO. - NÚMERO DE RUC:
[RUC]

QUINTO. - DATOS DEL DEMANDADO:
Nombres y apellidos: [DEMANDADO_NOMBRES_APELLIDOS]
Cédula: [DEMANDADO_CEDULA]
Nacionalidad: [DEMANDADO_NACIONALIDAD]
Profesión: [DEMANDADO_PROFESION]
Edad: [DEMANDADO_EDAD]
Provincia: [DEMANDADO_PROVINCIA]
Cantón: [DEMANDADO_CANTON]
Calle primaria: [DEMANDADO_CALLE_PRIMARIA]
Calle secundaria: [DEMANDADO_CALLE_SECUNDARIA]
Número de casa: [DEMANDADO_NUMERO_CASA]
Descripción de la vivienda: [DEMANDADO_DESCRIPCION_VIVIENDA]
Dirección electrónica: [DEMANDADO_DIR_ELECTRONICA]

SEXTO. - NARRACIÓN DETALLADA Y NUMERADA DE LOS HECHOS:
[HECHOS]

SÉPTIMO. - FUNDAMENTOS DE DERECHO:
[FUNDAMENTOS_DERECHO]

OCTAVO. - SOLICITUD DE ACCESO JUDICIAL A PRUEBAS:
[ACCESO_PRUEBAS]

NOVENO. - PRETENSIÓN:
[PRETENSION]

DÉCIMO. - CUANTÍA DEL PROCESO:
[CUANTIA]

DÉCIMO PRIMERO. - PROCEDIMIENTO QUE DEBE SEGUIRSE:
[PROCEDIMIENTO]

DÉCIMO SEGUNDO. - FIRMAS DEL ACTOR Y ABOGADO:
[FIRMAS]

DÉCIMO TERCERO. - OTROS REQUISITOS PERTINENTES AL CASO:
[OTROS]
"""

# Mapeo de ordinales para numerar secciones
ORDINALES = {
    1: "PRIMERO",
    2: "SEGUNDO",
    3: "TERCERO",
    4: "CUARTO",
    5: "QUINTO",
    6: "SEXTO",
    7: "SÉPTIMO",
    8: "OCTAVO",
    9: "NOVENO",
    10: "DÉCIMO",
    11: "DÉCIMO PRIMERO",
    12: "DÉCIMO SEGUNDO",
    13: "DÉCIMO TERCERO",
    14: "DÉCIMO CUARTO",
    15: "DÉCIMO QUINTO",
}


def _ordinal(n):
    """Devuelve la palabra ordinal correspondiente a ``n``."""
    return ORDINALES.get(n, str(n))

# Lista global con todas las secciones que forman el formulario de demanda
SECCIONES_DEMANDA = [
    "DESIGNACION_JUZGADOR",
    "DATOS_ACTOR",
    "DATOS_DEFENSOR",
    "RUC",
    "DATOS_DEMANDADO",
    "HECHOS",
    "FUNDAMENTOS_DERECHO",
    "ACCESO_PRUEBAS",
    "PRETENSION",
    "CUANTIA",
    "PROCEDIMIENTO",
    "FIRMAS",
    "OTROS",
]

# Campos detallados para la sección "DATOS_ACTOR"
CAMPOS_DATOS_ACTOR = [
    "ACTOR_NOMBRES_APELLIDOS",
    "ACTOR_CEDULA",
    "ACTOR_PASAPORTE",
    "ACTOR_ESTADO_CIVIL",
    "ACTOR_EDAD",
    "ACTOR_PROFESION",
    "ACTOR_PROVINCIA",
    "ACTOR_CANTON",
    "ACTOR_CALLE_PRIMARIA",
    "ACTOR_CALLE_SECUNDARIA",
    "ACTOR_NUMERO_CASA",
    "ACTOR_DIR_ELECTRONICA",
]

# Campos detallados para la sección "DATOS_DEFENSOR"
CAMPOS_DATOS_DEFENSOR = [
    "DEFENSOR_NOMBRE",
    "CASILLERO_JUDICIAL",
    "REPRESENTA_COMO_PROCURADOR",
    "DATOS_REPRESENTADO",
]

# Campos detallados para la sección "DATOS_DEMANDADO"
CAMPOS_DATOS_DEMANDADO = [
    "DEMANDADO_NOMBRES_APELLIDOS",
    "DEMANDADO_CEDULA",
    "DEMANDADO_NACIONALIDAD",
    "DEMANDADO_PROFESION",
    "DEMANDADO_EDAD",
    "DEMANDADO_PROVINCIA",
    "DEMANDADO_CANTON",
    "DEMANDADO_CALLE_PRIMARIA",
    "DEMANDADO_CALLE_SECUNDARIA",
    "DEMANDADO_NUMERO_CASA",
    "DEMANDADO_DESCRIPCION_VIVIENDA",
    "DEMANDADO_DIR_ELECTRONICA",
]


# ---------------------------------------------------------------------------
# Utilidades de PDF y vectores
# ---------------------------------------------------------------------------

def load_and_chunk_pdfs_from_folder(folder_path):
    """Carga PDFs y TXTs de una carpeta y los divide en fragmentos."""
    docs = []
    if not os.path.isdir(folder_path):
        return []

    doc_type_label = os.path.basename(folder_path.rstrip(os.sep))
    pdfs = glob.glob(os.path.join(folder_path, "**", "*.pdf"), recursive=True)
    txts = glob.glob(os.path.join(folder_path, "**", "*.txt"), recursive=True)

    for pdf_path in pdfs:
        loader = PyPDFLoader(pdf_path, extract_images=False)
        cargados = []
        try:
            cargados = loader.load()
        except PermissionError:
            print(
                f"⚠️ Permiso denegado al leer '{pdf_path}'. "
                "Revisa si está abierto o los permisos de archivo."
            )
        except Exception as e:
            # Si hay un error al leer, se intentará con OCR
            print(f"Error cargando '{pdf_path}': {e}")
            cargados = []

        text_available = any(d.page_content.strip() for d in cargados)
        if text_available:
            for d in cargados:
                d.metadata["doc_type"] = doc_type_label
                d.metadata["source"] = os.path.basename(d.metadata["source"])
            docs.extend(cargados)
        else:
            try:
                texto_ocr = read_pdf_text(pdf_path)
            except Exception:
                texto_ocr = ""
            if texto_ocr.strip():
                docs.append(
                    Document(
                        page_content=texto_ocr,
                        metadata={
                            "doc_type": doc_type_label,
                            "source": os.path.basename(pdf_path),
                        },
                    )
                )

    for txt_path in txts:
        try:
            with open(txt_path, "r", encoding="utf-8") as fh:
                contenido = fh.read()
            if contenido.strip():
                docs.append(
                    Document(
                        page_content=contenido,
                        metadata={
                            "doc_type": doc_type_label,
                            "source": os.path.basename(txt_path),
                        },
                    )
                )
        except PermissionError:
            print(
                f"⚠️ Permiso denegado al leer '{txt_path}'. "
                "Revisa si está abierto o los permisos de archivo."
            )
        except Exception as e:
            print(f"Error cargando '{txt_path}': {e}")

    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_documents(docs)


def build_or_load_vectorstore(path_docs, path_vector, force_rebuild=False, extra_docs=None):
    """Construye o carga un vectorstore para los documentos de `path_docs`.

    Si ``force_rebuild`` es ``True`` se elimina la carpeta ``path_vector``
    antes de crear un nuevo índice.
    """
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

    if force_rebuild and os.path.isdir(path_vector):
        try:
            shutil.rmtree(path_vector, ignore_errors=True)
        except PermissionError:
            print(
                f"⚠️ Permiso denegado al modificar '{path_vector}'. "
                "Revisa si está abierto o los permisos de archivo."
            )

    if os.path.isdir(path_vector) and os.listdir(path_vector):
        try:
            vs = FAISS.load_local(
                path_vector,
                embeddings,
                allow_dangerous_deserialization=True,
            )
            if extra_docs:
                vs.add_documents(extra_docs)
                vs.save_local(path_vector)
            return vs
        except PermissionError:
            print(
                f"⚠️ Permiso denegado al leer índice '{path_vector}'. "
                "Revisa si está abierto o los permisos de archivo."
            )

    chunks = []
    if os.path.isdir(path_docs):
        for sub in glob.glob(os.path.join(path_docs, "*")):
            if os.path.isdir(sub):
                chunks.extend(load_and_chunk_pdfs_from_folder(sub))
        chunks.extend(load_and_chunk_pdfs_from_folder(path_docs))
    else:
        raise ValueError(f"❌ La carpeta '{path_docs}' no existe o no es un directorio.")

    if extra_docs:
        chunks.extend(extra_docs)

    if not chunks:
        raise ValueError(f"❌ No se encontraron documentos PDF en '{path_docs}'.")
    vs = FAISS.from_documents(chunks, embeddings)
    try:
        vs.save_local(path_vector)
    except PermissionError:
        print(
            f"⚠️ Permiso denegado al guardar índice en '{path_vector}'. "
            "Revisa si está abierto o los permisos de archivo."
        )
    return vs


def load_email_docs(nombre_caso):
    """Return Document objects for emails associated with the case."""
    caso_id = db.get_caso_id_by_nombre(nombre_caso)
    if caso_id is None:
        return []
    docs = []
    for d in db.get_email_documents(caso_id):
        d.metadata["doc_type"] = "email"
        docs.append(d)
    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_documents(docs)


def _decode_header(value):
    """Return a decoded header value."""
    if not value:
        return ""
    decoded = decode_header(value)
    parts = []
    for text, enc in decoded:
        if isinstance(text, bytes):
            parts.append(text.decode(enc or "utf-8", errors="replace"))
        else:
            parts.append(text)
    return "".join(parts)


def get_pdf_attachments(msg):
    """Return list of (filename, bytes) for PDF attachments in the email."""
    attachments = []
    if not msg.is_multipart():
        return attachments
    for part in msg.walk():
        filename = part.get_filename()
        if not filename:
            continue
        filename = _decode_header(filename)
        ctype = part.get_content_type()
        if ctype == "application/pdf" or filename.lower().endswith(".pdf"):
            data = part.get_payload(decode=True)
            if data:
                attachments.append((filename, data))
    return attachments


def cargar_textos_demandas(carpeta):
    """Devuelve un diccionario ruta_relativa -> texto"""
    textos = {}
    for pdf in glob.glob(os.path.join(carpeta, "**", "*.pdf"), recursive=True):
        try:
            loader = PyPDFLoader(pdf, extract_images=False)
            docs = loader.load()
            rel = os.path.relpath(pdf, carpeta)
            textos[rel] = "\n".join(d.page_content for d in docs)
        except Exception as e:
            print(f"Error leyendo {pdf}: {e}")
    return textos


def _limpiar_texto(texto):
    """Normaliza espacios y saltos de línea para una lectura más clara."""
    texto = re.sub(r"[ \t]+\n", "\n", texto)
    texto = re.sub(r"\n[ \t]+", "\n", texto)
    texto = re.sub(r"[ \t]{2,}", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def _ocr_pdf(path: str) -> str:
    """Return OCR extracted text from the PDF if tesseract is available."""
    if convert_from_path is None or pytesseract is None:
        return ""
    try:
        images = convert_from_path(path)
    except Exception:
        return ""
    parts: list[str] = []
    for img in images:
        try:
            parts.append(pytesseract.image_to_string(img, lang="spa"))
        except Exception:
            continue
    return _limpiar_texto("\n".join(parts))


def _es_pdf_escaneado(path: str) -> bool:
    """Return True when the PDF lacks extractable text."""
    try:
        from pdfminer.high_level import extract_text
        texto = extract_text(path)
    except Exception:
        try:
            loader = PyPDFLoader(path, extract_images=False)
            docs = loader.load()
            texto = "\n".join(d.page_content for d in docs)
        except Exception:
            return True
    texto = _limpiar_texto(texto or "")
    return not bool(texto)


def leer_texto_pdf(nombre_archivo, carpeta=DEMANDAS_DIR):
    """Devuelve el texto completo de un PDF, usando OCR si es necesario.

    Si el texto ya fue extraído previamente se lee de ``<nombre>.txt`` en la
    misma carpeta para evitar repetir la operación de OCR.
    """
    path = os.path.join(carpeta, nombre_archivo)
    if not os.path.isfile(path):
        return ""

    cache_path = os.path.splitext(path)[0] + ".txt"
    if os.path.isfile(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as fh:
                return fh.read()
        except Exception:
            pass

    texto = ""
    if not _es_pdf_escaneado(path):
        try:
            loader = PyPDFLoader(path, extract_images=False)
            docs = loader.load()
            texto = _limpiar_texto("\n".join(d.page_content for d in docs))
        except Exception:
            texto = ""

    if not texto:
        texto = _ocr_pdf(path)

    if texto:
        try:
            with open(cache_path, "w", encoding="utf-8") as fh:
                fh.write(texto)
        except Exception:
            pass

    return texto


def _slugify(texto):
    """Devuelve una versi\xC3\B3n simplificada para usar como identificador."""
    texto = texto.lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = texto.encode("ascii", "ignore").decode("ascii")
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_")


def parsear_plantilla_desde_pdf(nombre_archivo, carpeta=DEMANDAS_DIR):
    """Lee un PDF y devuelve su texto sin buscar placeholders."""
    path = os.path.join(carpeta, nombre_archivo)
    if not os.path.isfile(path):
        return "", []

    if _es_pdf_escaneado(path):
        return _ocr_pdf(path), []

    try:
        from pdfminer.high_level import extract_text
        texto_original = extract_text(path)
        texto_original = _limpiar_texto(texto_original)
    except Exception:
        texto_original = ""
    if not texto_original:
        texto_original = _ocr_pdf(path)

    return texto_original, []


def extraer_datos_de_texto(texto):
    """Extrae datos básicos del texto de una demanda.

    Intenta identificar los datos del actor y del demandado buscando las
    secciones correspondientes en el escrito. Si no se encuentra una estructura
    reconocible, se mantiene una extracción simple de nombre, cédula y edad
    sobre el texto completo como en versiones anteriores.
    """

    datos = {}

    # Primero, intentar extraer datos dentro de las secciones de la demanda
    secciones = extraer_secciones_demanda(texto)

    def _extraer_persona(seccion):
        info = {}
        m = re.search(r"([A-ZÁÉÍÓÚÑ ]+),", seccion)
        if m:
            info["NOMBRES_APELLIDOS"] = m.group(1).title()
        m = re.search(r"c[ée]dula(?:\s+de\s+ciudadanía)?[^0-9]*(\d[\d-]*)", seccion, re.IGNORECASE)
        if m:
            info["CEDULA"] = m.group(1)
        m = re.search(r"(\d{1,3})\s*(?:años|anos)\s+de\s+edad", seccion, re.IGNORECASE)
        if m:
            info["EDAD"] = m.group(1)
        return info

    if secciones:
        if "DATOS_ACTOR" in secciones:
            info = _extraer_persona(secciones["DATOS_ACTOR"])
            if "NOMBRES_APELLIDOS" in info:
                datos["ACTOR_NOMBRES_APELLIDOS"] = info["NOMBRES_APELLIDOS"]
                datos.setdefault("NOMBRE", info["NOMBRES_APELLIDOS"])
            if "CEDULA" in info:
                datos["ACTOR_CEDULA"] = info["CEDULA"]
                datos.setdefault("CEDULA", info["CEDULA"])
            if "EDAD" in info:
                datos["ACTOR_EDAD"] = info["EDAD"]
                datos.setdefault("EDAD", info["EDAD"])

        if "DATOS_DEMANDADO" in secciones:
            info = _extraer_persona(secciones["DATOS_DEMANDADO"])
            if "NOMBRES_APELLIDOS" in info:
                datos["DEMANDADO_NOMBRES_APELLIDOS"] = info["NOMBRES_APELLIDOS"]
            if "CEDULA" in info:
                datos["DEMANDADO_CEDULA"] = info["CEDULA"]
            if "EDAD" in info:
                datos["DEMANDADO_EDAD"] = info["EDAD"]

    # Fallback: extracción general si aún faltan datos básicos
    if "NOMBRE" not in datos:
        m = re.search(r"([A-ZÁÉÍÓÚÑ ]+),\s+de\s+nacionalidad", texto)
        if m:
            datos["NOMBRE"] = m.group(1).title()
    if "CEDULA" not in datos:
        m = re.search(r"c[ée]dula(?:\s+de\s+ciudadanía)?[^0-9]*(\d[\d-]*)", texto)
        if m:
            datos["CEDULA"] = m.group(1)
    if "EDAD" not in datos:
        m = re.search(r"(\d{1,3})\s*(?:años|anos)\s+de\s+edad", texto, re.IGNORECASE)
        if m:
            datos["EDAD"] = m.group(1)

    return datos


ENTIDADES_FILE = Path(__file__).resolve().parents[1] / "data" / "entidades.json"
try:
    with open(ENTIDADES_FILE, "r", encoding="utf8") as fh:
        ENTIDADES_PREDEFINIDAS = json.load(fh)
except Exception:  # noqa: BLE001 - si no existe el archivo, no se detectan entidades
    ENTIDADES_PREDEFINIDAS = []

ENTIDADES_ENCONTRADAS = (
    Path(__file__).resolve().parents[1] / "data" / "entidades_encontradas.json"
)


def extraer_entidades_de_texto(texto):
    """Devuelve las entidades predefinidas presentes como ``{{ENTIDAD}}`` en ``texto``."""
    encontrados: list[str] = []
    for nombre in ENTIDADES_PREDEFINIDAS:
        marcador = f"{{{{{nombre}}}}}"
        if marcador in texto:
            encontrados.append(nombre)
    return encontrados


def registrar_entidades_encontradas(texto, path: Path | None = None):
    """Registra en ``path`` las entidades detectadas en ``texto`` y las devuelve."""
    if path is None:
        path = ENTIDADES_ENCONTRADAS
    encontrados = extraer_entidades_de_texto(texto)
    if not encontrados:
        return []
    try:
        with open(path, "r", encoding="utf8") as fh:
            actuales = set(json.load(fh))
    except Exception:
        actuales = set()
    nuevos = sorted(actuales.union(encontrados))
    if nuevos != list(actuales):
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf8") as fh:
            json.dump(nuevos, fh, ensure_ascii=False, indent=2)
    return encontrados


def registrar_entidades_por_llm(texto, llm_model=None, path: Path | None = None):
    """Detecta entidades mediante un LLM y las registra en ``path``.

    ``texto`` es el contenido de la demanda donde se buscarán los
    placeholders. Solo se registrarán aquellas entidades que coincidan con
    las predefinidas en ``ENTIDADES_PREDEFINIDAS``. El resultado es la lista
    de entidades detectadas y registradas. Si no se encuentra ninguna,
    devuelve una lista vacía.
    """

    if path is None:
        path = ENTIDADES_ENCONTRADAS

    encontrados = extraer_entidades_por_llm(texto, llm_model)
    if ENTIDADES_PREDEFINIDAS:
        encontrados = [e for e in encontrados if e in ENTIDADES_PREDEFINIDAS]
    if not encontrados:
        return []

    try:
        with open(path, "r", encoding="utf8") as fh:
            actuales = set(json.load(fh))
    except Exception:
        actuales = set()

    nuevos = sorted(actuales.union(encontrados))
    if nuevos != list(actuales):
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf8") as fh:
            json.dump(nuevos, fh, ensure_ascii=False, indent=2)

    return encontrados


def reemplazar_puntos(texto, valores):
    """Reemplaza secuencias de puntos o guiones bajos con los valores dados.

    Los grupos de tres o más ``.`` o ``______`` (seis o más ``_``) se sustituyen uno a uno por los
    elementos de ``valores`` en el orden provisto. Si hay más grupos que
    valores, los restantes se dejan sin modificar.
    """

    iter_vals = iter(valores)

    def _reemplazo(match):
        try:
            return str(next(iter_vals))
        except StopIteration:
            return match.group(0)

    return re.sub(r"(?:\.{3,}|_{5,})", _reemplazo, texto)


def extraer_secciones_demanda(texto):
    """Devuelve un mapeo de secciones del texto de la demanda.

    Retorna un diccionario donde las claves son los nombres definidos en
    ``SECCIONES_DEMANDA`` y los valores son los fragmentos de ``texto`` que
    corresponden a cada ordinal enumerado en el formato de la demanda.
    """

    # Construir un regex que detecte los encabezados del tipo
    # "PRIMERO. -", "DÉCIMO TERCERO. -", etc.
    ordinal_to_key = {
        ORDINALES[i]: SECCIONES_DEMANDA[i - 1]
        for i in range(1, min(len(SECCIONES_DEMANDA), len(ORDINALES)) + 1)
    }

    pattern = re.compile(
        r'^(%s)\.\s*-.*?:\s*' % "|".join(map(re.escape, ordinal_to_key)),
        re.MULTILINE,
    )

    partes = pattern.split(texto)

    secciones = {}
    for i in range(1, len(partes), 2):
        ordinal = partes[i]
        contenido = partes[i + 1].strip()
        clave = ordinal_to_key.get(ordinal)
        if clave:
            secciones[clave] = contenido

    return secciones


# Cache en memoria para resultados de demandas procesadas
_DEM_J_CACHE: dict[str, dict] = {}


def extraer_json_demanda(texto_o_ruta: str) -> dict:
    """Extrae secciones y campos de una demanda y devuelve un JSON (dict).

    ``texto_o_ruta`` puede ser texto plano o la ruta a un PDF. Si se entrega
    un archivo, se calcula su hash para reutilizar un resultado previo almacenado
    en memoria o en disco (``<archivo>.json``)."""

    texto = texto_o_ruta
    cache_key = None
    pdf_path = None

    if os.path.isfile(texto_o_ruta):
        pdf_path = texto_o_ruta
        try:
            cache_key = hash_for_pdf(pdf_path)
        except Exception:
            cache_key = None

        if cache_key and cache_key in _DEM_J_CACHE:
            return _DEM_J_CACHE[cache_key]

        cache_file = f"{pdf_path}.json"
        if cache_key and os.path.isfile(cache_file):
            try:
                with open(cache_file, "r", encoding="utf-8") as fh:
                    data = json.load(fh)
                if data.get("_hash") == cache_key:
                    _DEM_J_CACHE[cache_key] = data["data"]
                    return data["data"]
            except Exception:
                pass

        texto = leer_texto_pdf(os.path.basename(pdf_path), os.path.dirname(pdf_path))

    secciones_raw = extraer_secciones_demanda(texto)
    resultado: dict[str, dict | str] = {}
    for sec, contenido in secciones_raw.items():
        campos = {}
        for linea in contenido.splitlines():
            if ":" in linea:
                k, v = linea.split(":", 1)
                campos[k.strip()] = v.strip()
        resultado[sec] = campos if campos else contenido

    if cache_key:
        _DEM_J_CACHE[cache_key] = resultado
        try:
            with open(f"{pdf_path}.json", "w", encoding="utf-8") as fh:
                json.dump({"_hash": cache_key, "data": resultado}, fh, ensure_ascii=False, indent=2)
        except Exception:
            pass

    return resultado


def extraer_entidades_por_llm(texto, llm_model=None):
    """Utiliza un LLM para identificar los placeholders requeridos en una demanda."""
    llm_model = llm_model or get_llm()
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Devuelve un JSON con un objeto 'fields' cuyas claves son los nombres "
                "de los campos faltantes en la demanda. Responde solo con JSON válido",
            ),
            ("human", "{texto}"),
        ]
    )
    chain = prompt | llm_model | StrOutputParser()
    try:
        salida = chain.invoke({"texto": texto})
        match = re.search(r"\{.*\}", salida, re.DOTALL)
        if not match:
            return []
        datos = json.loads(match.group(0))
        fields = datos.get("fields")
        if isinstance(fields, dict):
            return list(fields.keys())
    except Exception:
        pass
    return []

def extraer_datos_por_llm(texto, llm_model=None):
    """Utiliza un LLM para extraer todas las entidades relevantes del texto."""
    llm_model = llm_model or get_llm()
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Extrae todas las entidades relevantes del siguiente texto. "
                "Devuelve únicamente un objeto JSON donde cada clave sea el nombre "
                "de la entidad en mayúsculas y guiones bajos, y cada valor el texto "
                "exacto encontrado. Omite las claves sin valor.",
            ),
            ("human", "{texto}"),
        ]
    )
    chain = prompt | llm_model | StrOutputParser()
    try:
        salida = chain.invoke({"texto": texto})
        datos = json.loads(salida)
        return {k.upper(): str(v) for k, v in datos.items()}
    except Exception:
        return {}


def extraer_datos_de_carpeta(carpeta):
    """Devuelve un diccionario con los datos extraídos de PDFs y TXTs."""
    datos = {}
    for pdf in glob.glob(os.path.join(carpeta, "**", "*.pdf"), recursive=True):
        texto = ""
        try:
            if not _es_pdf_escaneado(pdf):
                loader = PyPDFLoader(pdf, extract_images=False)
                docs = loader.load()
                texto = _limpiar_texto("\n".join(d.page_content for d in docs))
        except Exception:
            texto = ""

        if not texto:
            texto = _ocr_pdf(pdf)

        if texto:
            info = extraer_datos_por_llm(texto)
            if not info:
                info = extraer_datos_de_texto(texto)
            for k, v in info.items():
                datos.setdefault(k, v)

    for txt in glob.glob(os.path.join(carpeta, "**", "*.txt"), recursive=True):
        try:
            with open(txt, "r", encoding="utf-8") as fh:
                texto = fh.read()
            info = extraer_datos_por_llm(texto)
            if not info:
                info = extraer_datos_de_texto(texto)
            for k, v in info.items():
                datos.setdefault(k, v)
        except Exception:
            continue

    return datos


def _info_json_path(nombre_caso):
    """Devuelve la ruta del archivo JSON con los datos del caso."""
    return os.path.join(CASOS_DIR_ROOT, nombre_caso, "info.json")


def cargar_info_caso(nombre_caso):
    """Carga los datos del caso desde ``info.json`` si existe."""
    path = _info_json_path(nombre_caso)
    if os.path.isfile(path):
        try:
            with open(path, "r", encoding="utf-8") as fh:
                return json.load(fh)
        except Exception:
            return None
    return None


def guardar_info_caso(nombre_caso, datos):
    """Guarda un diccionario con datos del caso en ``info.json``."""
    path = _info_json_path(nombre_caso)
    try:
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(datos, fh, indent=2, ensure_ascii=False)
    except Exception:
        pass


def analizar_caso(nombre_caso, ctx=None):
    """Carga o genera los datos del caso y asegura su vectorstore."""
    if ctx is None:
        ctx = default_context
    nombre = nombre_caso.strip()
    if not nombre:
        return

    datos = cargar_info_caso(nombre)
    carpeta_caso = os.path.join(CASOS_DIR_ROOT, nombre)

    if datos is None:
        datos = extraer_datos_de_carpeta(carpeta_caso)
        guardar_info_caso(nombre, datos)

    ctx.datos_basicos_casos[nombre] = datos or {}

    if nombre not in ctx.vectorstores_por_caso:
        carpeta_index = os.path.join(VECTOR_DB_CASOS, nombre)
        email_docs = load_email_docs(nombre)
        ctx.vectorstores_por_caso[nombre] = build_or_load_vectorstore(
            carpeta_caso, carpeta_index, extra_docs=email_docs
        )

    if nombre not in ctx.memories_por_caso:
        ctx.memories_por_caso[nombre] = ChatMessageHistory()



def reemplazar_datos(texto, info_origen, info_destino):
    """Reemplaza en `texto` los valores de info_origen por los de info_destino."""
    for clave in ("NOMBRE", "CEDULA", "EDAD"):
        if clave in info_origen and clave in info_destino:
            texto = re.sub(re.escape(info_origen[clave]), info_destino[clave], texto, flags=re.IGNORECASE)
    return texto



# ---------------------------------------------------------------------------
# Persistencia de configuración
# ---------------------------------------------------------------------------

def cargar_config():
    if os.path.isfile(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def guardar_config(config_dict):
    try:
        os.makedirs(os.path.dirname(CONFIG_PATH), exist_ok=True)
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(config_dict, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"⚠️ Error guardando config.json: {e}")


class DemandasContext:
    """Objeto que almacena el estado compartido de la aplicación."""

    def __init__(self):
        self.juris_vectorstore = None
        self.juris_index = None
        self.vectorstores_por_caso = {}
        self.memories_por_caso = {}
        self.demandas_vectorstore = None
        self.legal_vectorstore = None
        self.memory_ejemplos = ChatMessageHistory()
        self.memory_palabras = ChatMessageHistory()

        self.last_generated_document = ""
        self.pending_placeholders = []
        self.placeholder_values = {}
        self.current_placeholder_index = 0
        self.partial_document = ""

        # Almacena el texto completo de cada plantilla de demanda cargada
        self.demandas_textos = {}

        self.datos_basicos_casos = {}
        self.datos_basicos_demandas = {}

        self.config_global = cargar_config()
        self.llm_provider = self.config_global.get("llm_provider", "openai")
        self.ruta_juris = _resolve_path(
            self.config_global.get("juris_path", ""), JURIS_DIR
        )
        self.ruta_casos_root = _resolve_path(
            self.config_global.get("casos_path", ""), CASOS_DIR_ROOT
        )
        self.ruta_demandas = _resolve_path(
            self.config_global.get("demandas_path", ""), DEMANDAS_DIR
        )
        self.ruta_areas_root = _resolve_path(
            self.config_global.get("areas_path", ""), AREAS_DIR_ROOT
        )
        os.makedirs(self.ruta_areas_root, exist_ok=True)
        self.ruta_legal_corpus = _resolve_path(
            self.config_global.get("legal_corpus_path", ""), LEGAL_CORPUS_DIR
        )


default_context = DemandasContext()

# Expose some attributes at module level for backwards compatibility
demandas_textos = default_context.demandas_textos
datos_basicos_casos = default_context.datos_basicos_casos
datos_basicos_demandas = default_context.datos_basicos_demandas


def resumir_caso(nombre_caso, ctx=default_context):
    """Devuelve un resumen breve del caso usando su vectorstore.

    Si el resumen ya existe en ``ctx.datos_basicos_casos`` se devuelve
    directamente para evitar llamadas innecesarias al LLM.
    """

    info = ctx.datos_basicos_casos.get(nombre_caso)
    if info and info.get("RESUMEN"):
        return info["RESUMEN"]

    if nombre_caso not in ctx.vectorstores_por_caso:
        return ""

    vs = ctx.vectorstores_por_caso[nombre_caso]
    history = ChatMessageHistory()
    raw = ConversationalRetrievalChain.from_llm(
        llm=get_llm(),
        retriever=vs.as_retriever(search_kwargs={"k": 15}),
        combine_docs_chain_kwargs={"prompt": summary_prompt},
    )
    chain = RunnableWithMessageHistory(
        raw,
        lambda _sid: history,
        input_messages_key="question",
        history_messages_key="chat_history",
    )
    pregunta = (
        "Resume brevemente los hechos relevantes del caso e identifica "
        "si es posible datos como nombres y cédulas de las partes."
    )
    respuesta = chain.invoke(
        {"question": pregunta},
        config={"configurable": {"session_id": "summary"}},
    )["answer"].strip()

    if info is None:
        info = {}
    info["RESUMEN"] = respuesta
    ctx.datos_basicos_casos[nombre_caso] = info
    guardar_info_caso(nombre_caso, info)

    return respuesta


def _resumir_texto_llm(texto, llm_model=None):
    """Devuelve un resumen breve para el texto dado usando un LLM."""
    if not texto.strip():
        return ""
    llm_model = llm_model or get_llm()
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Eres un abogado virtual de Ecuador. Resume brevemente el siguiente texto en español.",
            ),
            ("human", "{texto}"),
        ]
    )
    chain = prompt | llm_model | StrOutputParser()
    try:
        return chain.invoke({"texto": texto}).strip()
    except Exception:
        return ""


def resumir_archivo_caso(nombre_caso, nombre_archivo, ctx=default_context, llm_model=None):
    """Devuelve un resumen para un documento específico de un caso."""
    nombre_caso = nombre_caso.strip()
    nombre_archivo = nombre_archivo.strip()
    if not nombre_caso or not nombre_archivo:
        return ""

    carpeta = os.path.join(CASOS_DIR_ROOT, nombre_caso)
    ext = os.path.splitext(nombre_archivo)[1].lower()
    texto = ""
    if ext == ".pdf":
        texto = leer_texto_pdf(nombre_archivo, carpeta=carpeta)
    else:
        path = os.path.join(carpeta, nombre_archivo)
        if os.path.isfile(path):
            try:
                with open(path, "r", encoding="utf-8") as fh:
                    texto = fh.read()
            except Exception:
                texto = ""

    return _resumir_texto_llm(texto, llm_model)


# ---------------------------------------------------------------------------
# Estado global (mantenido en un objeto de contexto)
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Lógica de generación de demandas
# ---------------------------------------------------------------------------


def chat_fn(
    mensaje,
    caso_seleccionado,
    ctx=default_context,
    extra_retriever=None,
    usar_jurisprudencia=True,
):
    """Atiende a un mensaje del usuario."""

    texto = mensaje.strip()
    low = texto.lower()

    if ctx.pending_placeholders:
        if low.startswith("omitir"):
            ctx.placeholder_values[ctx.pending_placeholders[ctx.current_placeholder_index]] = f"[{ctx.pending_placeholders[ctx.current_placeholder_index]}]"
        else:
            ctx.placeholder_values[ctx.pending_placeholders[ctx.current_placeholder_index]] = texto
        ctx.current_placeholder_index += 1
        if ctx.current_placeholder_index < len(ctx.pending_placeholders):
            return (
                f"Ingresa el valor para '{ctx.pending_placeholders[ctx.current_placeholder_index]}' "
                "o escribe 'omitir' para dejarlo en blanco."
            )
        else:
            updated = ctx.partial_document
            for ph, val in ctx.placeholder_values.items():
                updated = updated.replace(f'[{ph}]', val)
            ctx.last_generated_document = updated
            ctx.partial_document = updated
            ctx.pending_placeholders = []
            ctx.placeholder_values = {}
            ctx.current_placeholder_index = 0
            return updated

    retrievers = []
    if caso_seleccionado and caso_seleccionado in ctx.vectorstores_por_caso:
        retrievers.append(
            ctx.vectorstores_por_caso[caso_seleccionado].as_retriever(search_kwargs={"k": 15})
        )

    if usar_jurisprudencia and ctx.juris_vectorstore:
        retrievers.append(
            ctx.juris_vectorstore.as_retriever(search_kwargs={"k": 15})
        )

    if extra_retriever:
        retrievers.append(extra_retriever)

    if not retrievers:
        return "⚠️ Debes seleccionar al menos un documento para chatear."

    if len(retrievers) == 1:
        retriever = retrievers[0]
    else:
        retriever = EnsembleRetriever(retrievers=retrievers, weights=[1] * len(retrievers))

    if caso_seleccionado and caso_seleccionado in ctx.memories_por_caso:
        history = ctx.memories_por_caso[caso_seleccionado]
    else:
        if "_GLOBAL_" not in ctx.memories_por_caso:
            ctx.memories_por_caso["_GLOBAL_"] = ChatMessageHistory()
        history = ctx.memories_por_caso["_GLOBAL_"]

    raw_chain = ConversationalRetrievalChain.from_llm(
        llm=get_llm(),
        retriever=retriever,
        combine_docs_chain_kwargs={"prompt": prompt},
    )
    chain = RunnableWithMessageHistory(
        raw_chain,
        lambda _sid: history,
        input_messages_key="question",
        history_messages_key="chat_history",
    )
    session_id = caso_seleccionado or "_GLOBAL_"
    resultado = chain.invoke(
        {"question": texto},
        config={"configurable": {"session_id": session_id}},
    )["answer"]
    return resultado


def generar_demanda_de_tipo(tipo_demanda, caso_seleccionado, ctx=default_context):
    """Genera una demanda a partir de un tipo y un caso."""

    if not ctx.demandas_vectorstore:
        return "⚠️ Debes cargar primero las plantillas de demandas (pestaña “⚙️ Configuración”)."

    retr_dem = ctx.demandas_vectorstore.as_retriever(
        search_kwargs={"k": 25, "filter": {"doc_type": os.path.basename(DEMANDAS_DIR.rstrip(os.sep))}}
    )
    history_rag = ChatMessageHistory()
    raw_chain_rag = ConversationalRetrievalChain.from_llm(
        llm=get_llm(),
        retriever=retr_dem,
        combine_docs_chain_kwargs={"prompt": prompt},
    )
    chain_rag = RunnableWithMessageHistory(
        raw_chain_rag,
        lambda _sid: history_rag,
        input_messages_key="question",
        history_messages_key="chat_history",
    )
    rag_resp = chain_rag.invoke(
        {"question": f"Por favor, dame una plantilla de demanda para: {tipo_demanda}"},
        config={"configurable": {"session_id": "rag"}},
    )
    plantilla_texto = rag_resp["answer"].strip()

    info_ejemplo = extraer_datos_de_texto(plantilla_texto)
    info_caso = ctx.datos_basicos_casos.get(caso_seleccionado, {})
    plantilla_texto = reemplazar_datos(plantilla_texto, info_ejemplo, info_caso)

    if not plantilla_texto:
        return f"❌ No se encontró ninguna plantilla relacionada con '{tipo_demanda}'."

    placeholders = re.findall(r"\[([^\]]+)\]", plantilla_texto)

    valores_dict = {}

    if caso_seleccionado and caso_seleccionado in ctx.vectorstores_por_caso:
        vs_caso = ctx.vectorstores_por_caso[caso_seleccionado]
        for ph in placeholders:
            if ph.upper() in info_caso:
                valores_dict[ph] = info_caso[ph.upper()]
            else:
                pregunta = f"¿Cuál es el {ph.replace('_', ' ')} en este caso?"
                history_case = ChatMessageHistory()
                raw_chain_case = ConversationalRetrievalChain.from_llm(
                    llm=get_llm(),
                    retriever=vs_caso.as_retriever(search_kwargs={"k": 15}),
                    combine_docs_chain_kwargs={"prompt": prompt},
                )
                chain_case = RunnableWithMessageHistory(
                    raw_chain_case,
                    lambda _sid: history_case,
                    input_messages_key="question",
                    history_messages_key="chat_history",
                )
                resp_case = chain_case.invoke(
                    {"question": pregunta},
                    config={"configurable": {"session_id": "case"}},
                )["answer"].strip()
                if resp_case and len(resp_case) > 3 and "no" not in resp_case.lower():
                    valores_dict[ph] = resp_case
                else:
                    valores_dict[ph] = f"[{ph}]"
    else:
        for ph in placeholders:
            if ph.upper() in info_caso:
                valores_dict[ph] = info_caso[ph.upper()]
            else:
                valores_dict[ph] = f"[{ph}]"

    texto_final = plantilla_texto
    for ph_key, ph_val in valores_dict.items():
        marcador = f"[{ph_key}]"
        texto_final = texto_final.replace(marcador, ph_val)

    ctx.pending_placeholders = re.findall(r"\[([^\]]+)\]", texto_final)
    ctx.placeholder_values = {}
    ctx.current_placeholder_index = 0
    ctx.partial_document = texto_final

    ctx.last_generated_document = texto_final

    if ctx.pending_placeholders:
        faltantes = ", ".join(ctx.pending_placeholders)
        mensaje_extra = (
            f"\n\nFaltan datos para: {faltantes}. "
            f"Ingresa el valor para '{ctx.pending_placeholders[0]}' o escribe 'omitir' para dejarlo en blanco."
        )
        return texto_final + mensaje_extra
    return texto_final


def generar_demanda_desde_pdf(
    nombre_archivo,
    caso_seleccionado,
    datos=None,
    ctx=default_context,
    carpeta=None,
):
    """Genera una demanda usando la plantilla del PDF seleccionado."""

    if carpeta is None:
        carpeta = ctx.ruta_demandas or DEMANDAS_DIR

    plantilla_texto, _ = parsear_plantilla_desde_pdf(nombre_archivo, carpeta=carpeta)
    if not plantilla_texto:
        plantilla_texto = ctx.demandas_textos.get(nombre_archivo, "")
    if not plantilla_texto:
        return f"❌ Demanda '{nombre_archivo}' no encontrada."

    info_ejemplo = ctx.datos_basicos_demandas.get(
        nombre_archivo, extraer_datos_de_texto(plantilla_texto)
    )
    info_caso = ctx.datos_basicos_casos.get(caso_seleccionado, {})
    plantilla_texto = reemplazar_datos(plantilla_texto, info_ejemplo, info_caso)

    return completar_documento_con_caso(
        plantilla_texto,
        caso_seleccionado,
        datos=datos,
        ctx=ctx,
    )


def generar_demanda_cogep(ctx=default_context):
    """Devuelve una plantilla de demanda basada en el art. 142 del COGEP."""
    texto = COGEP_TEMPLATE.strip()
    ctx.pending_placeholders = re.findall(r"\[([^\]]+)\]", texto)
    ctx.placeholder_values = {}
    ctx.current_placeholder_index = 0
    ctx.partial_document = texto
    ctx.last_generated_document = texto

    if ctx.pending_placeholders:
        faltantes = ", ".join(ctx.pending_placeholders)
        mensaje_extra = (
            f"\n\nFaltan datos para: {faltantes}. "
            f"Ingresa el valor para '{ctx.pending_placeholders[0]}' o escribe 'omitir' para dejarlo en blanco."
        )
        return texto + mensaje_extra

    return texto


def obtener_dato_de_caso(placeholder, caso_seleccionado, ctx=default_context):
    """Devuelve un posible valor para ``placeholder`` usando datos del caso."""
    info_caso = ctx.datos_basicos_casos.get(caso_seleccionado, {})
    if placeholder.upper() in info_caso:
        return info_caso[placeholder.upper()]
    if (
        caso_seleccionado
        and caso_seleccionado in ctx.vectorstores_por_caso
    ):
        vs_caso = ctx.vectorstores_por_caso[caso_seleccionado]
        pregunta = f"¿Cuál es el {placeholder.replace('_', ' ')} en este caso?"
        history_case = ChatMessageHistory()
        raw_case = ConversationalRetrievalChain.from_llm(
            llm=get_llm(),
            retriever=vs_caso.as_retriever(search_kwargs={"k": 15}),
            combine_docs_chain_kwargs={"prompt": prompt},
        )
        chain_case = RunnableWithMessageHistory(
            raw_case,
            lambda _sid: history_case,
            input_messages_key="question",
            history_messages_key="chat_history",
        )
        resp = chain_case.invoke(
            {"question": pregunta},
            config={"configurable": {"session_id": "case"}},
        )["answer"].strip()
        if resp and len(resp) > 3 and "no" not in resp.lower():
            return resp
    return None


def completar_documento_con_caso(texto_base, caso_seleccionado, datos=None, ctx=default_context):
    """Reemplaza placeholders en ``texto_base`` usando ``datos`` y el caso."""

    placeholders = re.findall(r"\[([^\]]+)\]", texto_base)
    valores = {}
    datos = datos or {}
    for ph in placeholders:
        val = datos.get(ph) or datos.get(ph.upper())
        if not val:
            val = obtener_dato_de_caso(ph, caso_seleccionado, ctx)
        valores[ph] = val if val else f"[{ph}]"

    texto_final = texto_base
    for ph, val in valores.items():
        texto_final = texto_final.replace(f"[{ph}]", val)

    # Replace dotted or underscored blanks with the provided data values
    if datos:
        datos_vals = [str(v) for v in datos.values() if v]
        texto_final = reemplazar_puntos(texto_final, datos_vals)

    ctx.pending_placeholders = re.findall(r"\[([^\]]+)\]", texto_final)
    ctx.placeholder_values = {}
    ctx.current_placeholder_index = 0
    ctx.partial_document = texto_final
    ctx.last_generated_document = texto_final

    if ctx.pending_placeholders:
        faltantes = ", ".join(ctx.pending_placeholders)
        mensaje_extra = (
            f"\n\nFaltan datos para: {faltantes}. "
            f"Ingresa el valor para '{ctx.pending_placeholders[0]}' o escribe 'omitir' para dejarlo en blanco."
        )
        return texto_final + mensaje_extra
    return texto_final


def generar_demanda_cogep_con_datos(caso_seleccionado, ctx=default_context):
    """Genera la plantilla del COGEP completando datos desde el caso."""
    texto = COGEP_TEMPLATE.strip()
    return completar_documento_con_caso(texto, caso_seleccionado, ctx=ctx)


def buscar_ejemplo_fn(mensaje, ctx=default_context):

    if not ctx.demandas_vectorstore:
        return "⚠️ Debes cargar primero las plantillas de demandas (pestaña “⚙️ Configuración”)."

    texto = mensaje.strip()
    retr = ctx.demandas_vectorstore.as_retriever(
        search_kwargs={"k": 25, "filter": {"doc_type": os.path.basename(DEMANDAS_DIR.rstrip(os.sep))}}
    )
    raw_chain = ConversationalRetrievalChain.from_llm(
        llm=get_llm(),
        retriever=retr,
        combine_docs_chain_kwargs={"prompt": prompt},
    )
    chain = RunnableWithMessageHistory(
        raw_chain,
        lambda _sid: ctx.memory_ejemplos,
        input_messages_key="question",
        history_messages_key="chat_history",
    )
    respuesta = chain.invoke(
        {"question": texto},
        config={"configurable": {"session_id": "ejemplos"}},
    )["answer"]
    return respuesta


def buscar_palabras_clave_fn(mensaje, ctx=default_context):

    if not ctx.juris_vectorstore:
        return "No se ha cargado jurisprudencia."


    if re.search(r"art(?:[íi]culo|\.)?\s*\d+", mensaje, re.IGNORECASE):

        return buscar_palabras_clave_exacta_fn(mensaje, ctx)

    texto = mensaje.strip()
    retr = ctx.juris_vectorstore.as_retriever(search_kwargs={"k": 25})
    raw_chain = ConversationalRetrievalChain.from_llm(
        llm=get_llm(),
        retriever=retr,
        combine_docs_chain_kwargs={"prompt": prompt},
        return_source_documents=True,
    )
    chain = RunnableWithMessageHistory(
        raw_chain,
        lambda _sid: ctx.memory_palabras,
        input_messages_key="question",
        history_messages_key="chat_history",
    )
    resultado = chain.invoke(
        {"question": texto},
        config={"configurable": {"session_id": "palabras"}},
    )
    respuesta = resultado["answer"]
    enlaces = []
    vistos = set()
    for doc in resultado.get("source_documents", []):
        meta = getattr(doc, "metadata", {}) or {}
        base = os.path.splitext(os.path.basename(meta.get("source", "")))[0]
        if base and base not in vistos:
            ruta = os.path.join(ctx.ruta_juris, f"{base}.pdf")
            enlaces.append(f"- {base}: {ruta}")
            vistos.add(base)
    if enlaces:
        respuesta += "\n\nEnlaces:\n" + "\n".join(enlaces)
    return respuesta


# Sinónimos simples para coincidencias directas en búsquedas exactas.
SINONIMOS = {
    "abuso": ["agresion", "agresión"],
    "agresion": ["abuso", "agresión"],
    "agresión": ["abuso", "agresion"],
}


def _buscar_articulos_por_palabras(texto, ctx):
    """Busca artículos que contengan todas las palabras o sus sinónimos."""

    palabras = [p for p in re.findall(r"\w+", texto.lower()) if len(p) > 2]
    grupos = []
    for p in palabras:
        grupo = {p}
        grupo.update(SINONIMOS.get(p, []))
        grupos.append(grupo)

    resultados = []
    vistos = set()

    for entrada in ctx.juris_index:
        txt = entrada["text"].lower()
        if all(any(s in txt for s in grupo) for grupo in grupos):
            key = (
                entrada["metadata"].get("documento"),
                entrada["metadata"].get("articulo"),
            )
            if key not in vistos:
                resultados.append(entrada)
                vistos.add(key)

    if not resultados:
        return "❌ No se encontraron artículos para la consulta."

    partes = []
    for entrada in resultados:
        meta = entrada["metadata"]
        doc = meta.get("documento", "")
        enlace = os.path.join(ctx.ruta_juris, f"{doc}.pdf")
        partes.append(
            f"Artículo {meta.get('articulo')} de {doc}:\n{entrada['text'].strip()}\nEnlace: {enlace}"
        )
    return "\n\n".join(partes)


def buscar_palabras_clave_exacta_fn(mensaje, ctx=default_context):
    """Realiza una búsqueda exacta de un artículo usando metadatos.

    Si no se detecta un número de artículo, se realiza una búsqueda híbrida
    semántica y por palabras, devolviendo únicamente los artículos encontrados.
    """

    if not ctx.juris_vectorstore:
        return "No se ha cargado jurisprudencia."
    if not ctx.juris_index:
        ctx.juris_index = build_exact_index(ctx.juris_vectorstore)

    match_art = re.search(r"art(?:[íi]culo|\.)?\s*(\d+)", mensaje, re.IGNORECASE)
    if not match_art:
        return _buscar_articulos_por_palabras(mensaje.strip(), ctx)
    num_art = match_art.group(1)

    match_doc = re.search(
        rf"art(?:[íi]culo|\.)?\s*{num_art}\s+de\s+(?:la|el|los|las)\s+([\w\sáéíóúñÁÉÍÓÚÑ]+)",
        mensaje,
        re.IGNORECASE,
    )

    doc_name = match_doc.group(1).strip() if match_doc else ""

    resultados = search_article(num_art, doc_name, ctx.juris_index)
    if not resultados:
        return "❌ No se encontró el artículo solicitado."

    partes = []
    for entrada in resultados:
        doc = entrada["metadata"].get("documento", "")
        texto = entrada["text"].strip()
        enlace = os.path.join(ctx.ruta_juris, f"{doc}.pdf")
        partes.append(f"Artículo {num_art} de {doc}:\n{texto}\nEnlace: {enlace}")

    return "\n\n".join(partes)




# XML no permite los caracteres de control de los rangos C0 y C1, con
# excepci\xC3\B3n de TAB (0x09), LF (0x0A) y CR (0x0D). El texto extra\xC3\ADdo de
# los PDFs puede incluirlos, por lo que es necesario filtrarlos antes de
# crear elementos de ``python-docx``. Ampliamos la expresi\xC3\B3n regular para
# cubrir tambi\xC3\A9n el rango 0x7F-0x9F (conocido como controles C1).
_CTRL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _remove_control_chars(text: str) -> str:
    """Return ``text`` stripped of XML-invalid control characters."""
    return _CTRL_CHARS_RE.sub("", text)


def exportar_a_word(texto_para_exportar, path_destino=None):
    if not texto_para_exportar:
        messagebox.showwarning("Exportar a Word", "❌ No hay texto para exportar.")
        return

    try:
        from docx import Document as DocxDocument
    except ImportError:
        messagebox.showerror(
            "python-docx no instalado",
            "Debes instalar `python-docx` para exportar a Word:\n\npip install python-docx",
        )
        return

    doc = DocxDocument()
    clean_text = _remove_control_chars(texto_para_exportar)
    for linea in clean_text.split("\n"):
        doc.add_paragraph(linea)

    if not path_destino:
        path_destino = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
            initialfile="demanda_generada.docx",
        )
        if not path_destino:
            return

    try:
        doc.save(path_destino)
        messagebox.showinfo("Exportar a Word", f"✅ Documento guardado en:\n\n{path_destino}")
    except Exception as e:
        messagebox.showerror("Error al guardar archivo", f"❌ No se pudo guardar el archivo:\n{e}")


def exportar_a_pdf(texto_para_exportar, path_destino=None):
    if not texto_para_exportar:
        messagebox.showwarning("Exportar a PDF", "❌ No hay texto para exportar.")
        return

    try:
        from fpdf import FPDF
    except ImportError:
        messagebox.showerror(
            "fpdf no instalado",
            "Debes instalar `fpdf` para exportar a PDF:\n\npip install fpdf",
        )
        return

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    clean_text = _remove_control_chars(texto_para_exportar)
    for linea in clean_text.split("\n"):
        pdf.multi_cell(0, 10, txt=linea)

    if not path_destino:
        path_destino = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("Archivos PDF", "*.pdf")],
            initialfile="documento.pdf",
        )
        if not path_destino:
            return

    try:
        pdf.output(path_destino)
        messagebox.showinfo("Exportar a PDF", f"✅ Documento guardado en:\n\n{path_destino}")
    except Exception as e:
        messagebox.showerror("Error al guardar archivo", f"❌ No se pudo guardar el archivo:\n{e}")


def exportar_a_txt(texto_para_exportar, path_destino=None):
    """Guarda ``texto_para_exportar`` en un archivo .txt."""

    if not texto_para_exportar:
        messagebox.showwarning("Exportar a TXT", "❌ No hay texto para exportar.")
        return

    if not path_destino:
        path_destino = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Archivos de texto", "*.txt")],
            initialfile="documento.txt",
        )
        if not path_destino:
            return

    try:
        clean_text = _remove_control_chars(texto_para_exportar)
        with open(path_destino, "w", encoding="utf-8") as f:
            f.write(clean_text)
        messagebox.showinfo("Exportar a TXT", f"✅ Documento guardado en:\n\n{path_destino}")
    except Exception as e:
        messagebox.showerror("Error al guardar archivo", f"❌ No se pudo guardar el archivo:\n{e}")


# ---------------------------------------------------------------------------
# Base de datos de casos
# ---------------------------------------------------------------------------

def init_db(path=DB_PATH):
    conn = sqlite3.connect(path)
    conn.execute(
        "CREATE TABLE IF NOT EXISTS casos (id INTEGER PRIMARY KEY AUTOINCREMENT, tipo TEXT, resultado TEXT, cliente TEXT)"
    )
    conn.commit()
    conn.close()


def registrar_caso(tipo, resultado, cliente, path=DB_PATH):
    init_db(path)
    conn = sqlite3.connect(path)
    conn.execute(
        "INSERT INTO casos (tipo, resultado, cliente) VALUES (?, ?, ?)",
        (tipo, resultado, cliente),
    )
    conn.commit()
    conn.close()


def obtener_conteos(path=DB_PATH):
    init_db(path)
    conn = sqlite3.connect(path)
    cur = conn.cursor()
    conteos_tipo = {row[0]: row[1] for row in cur.execute("SELECT tipo, COUNT(*) FROM casos GROUP BY tipo")}
    conteos_resultado = {row[0]: row[1] for row in cur.execute("SELECT resultado, COUNT(*) FROM casos GROUP BY resultado")}
    conteos_cliente = {row[0]: row[1] for row in cur.execute("SELECT cliente, COUNT(*) FROM casos GROUP BY cliente")}
    conn.close()
    return {
        "por_tipo": conteos_tipo,
        "por_resultado": conteos_resultado,
        "por_cliente": conteos_cliente,
    }


# ---------------------------------------------------------------------------
# Funciones de gestión de datos
# ---------------------------------------------------------------------------

def guardar_jurisprudencia_por_carpeta(carpeta_sel, ctx=default_context):
    try:
        for f in glob.glob(os.path.join(JURIS_DIR, "*")):
            try:
                if os.path.isfile(f):
                    os.remove(f)
                else:
                    shutil.rmtree(f)
            except Exception:
                pass
        for root, _, files in os.walk(carpeta_sel):
            rel = os.path.relpath(root, carpeta_sel)
            dest_folder = os.path.join(JURIS_DIR, rel)
            os.makedirs(dest_folder, exist_ok=True)
            for fname in files:
                if fname.lower().endswith(".pdf"):
                    src = os.path.join(root, fname)
                    dest = os.path.join(dest_folder, fname)
                    try:
                        shutil.copy(src, dest)
                    except PermissionError:
                        print(f"⚠️ Permiso denegado al copiar '{src}'. Se omite.")

        ctx.juris_vectorstore = build_or_load_vectorstore(
            JURIS_DIR, VECTOR_DB_JURIS, force_rebuild=True
        )
        ctx.juris_index = build_exact_index(ctx.juris_vectorstore)
        ctx.ruta_juris = carpeta_sel
        ctx.config_global["juris_path"] = ctx.ruta_juris
        guardar_config(ctx.config_global)
        return "✅ Jurisprudencia cargada e indexada correctamente."
    except Exception as e:
        return f"❌ Error al cargar jurisprudencia: {e}"


def actualizar_jurisprudencia(ctx=default_context):
    src_path = _resolve_path(ctx.ruta_juris, "")
    if not src_path or not os.path.isdir(src_path):
        return "⚠️ No existe una ruta de jurisprudencia válida. Debes cargar primero."
    try:
        for f in glob.glob(os.path.join(JURIS_DIR, "*")):
            try:
                if os.path.isfile(f):
                    os.remove(f)
                else:
                    shutil.rmtree(f)
            except Exception:
                pass

        for root, _, files in os.walk(ctx.ruta_juris):
            rel = os.path.relpath(root, ctx.ruta_juris)

            dest_folder = os.path.join(JURIS_DIR, rel)
            os.makedirs(dest_folder, exist_ok=True)
            for fname in files:
                if fname.lower().endswith(".pdf"):
                    src = os.path.join(root, fname)
                    dest = os.path.join(dest_folder, fname)
                    try:
                        shutil.copy(src, dest)
                    except PermissionError:
                        print(f"⚠️ Permiso denegado al copiar '{src}'. Se omite.")

        ctx.juris_vectorstore = build_or_load_vectorstore(
            JURIS_DIR, VECTOR_DB_JURIS, force_rebuild=True
        )
        ctx.juris_index = build_exact_index(ctx.juris_vectorstore)
        return "✅ Jurisprudencia actualizada correctamente."
    except Exception as e:
        return f"❌ Error al actualizar jurisprudencia: {e}"


def eliminar_jurisprudencia(ctx=default_context):
    """Elimina la jurisprudencia cargada y su índice vectorial."""
    try:
        if os.path.isdir(JURIS_DIR):
            shutil.rmtree(JURIS_DIR)
        os.makedirs(JURIS_DIR, exist_ok=True)

        if os.path.isdir(VECTOR_DB_JURIS):
            shutil.rmtree(VECTOR_DB_JURIS)
        os.makedirs(VECTOR_DB_JURIS, exist_ok=True)

        ctx.juris_vectorstore = None
        ctx.juris_index = None
        ctx.ruta_juris = ""
        if ctx.config_global.get("juris_path"):
            ctx.config_global["juris_path"] = ""
            guardar_config(ctx.config_global)
        return "✅ Jurisprudencia eliminada correctamente."
    except Exception as e:
        return f"❌ Error al eliminar jurisprudencia: {e}"


def agregar_jurisprudencia_pdf(ruta_pdf, ctx=default_context):
    """Agrega un archivo PDF al directorio de jurisprudencia e indexa."""

    if not ruta_pdf.lower().endswith(".pdf") or not os.path.isfile(ruta_pdf):
        return "⚠️ Debes seleccionar un archivo PDF válido."

    try:
        os.makedirs(JURIS_DIR, exist_ok=True)
        dest = os.path.join(JURIS_DIR, os.path.basename(ruta_pdf))
        shutil.copy(ruta_pdf, dest)
        ctx.juris_vectorstore = build_or_load_vectorstore(
            JURIS_DIR, VECTOR_DB_JURIS, force_rebuild=True
        )
        ctx.juris_index = build_exact_index(ctx.juris_vectorstore)
        return "✅ Jurisprudencia agregada e indexada correctamente."
    except Exception as e:
        return f"❌ Error al agregar jurisprudencia: {e}"


def guardar_caso_por_carpeta(carpeta_sel, nombre_caso, ctx=default_context):

    if not nombre_caso.strip():
        return "⚠️ Debes escribir un nombre válido para la carpeta del caso.", []

    nombre_carpeta = nombre_caso.strip()
    carpeta_dest = os.path.join(CASOS_DIR_ROOT, nombre_carpeta)

    try:
        if os.path.isdir(carpeta_dest):
            shutil.rmtree(carpeta_dest)

        for root, _, files in os.walk(carpeta_sel):
            rel = os.path.relpath(root, carpeta_sel)
            dest_folder = os.path.join(carpeta_dest, rel)
            os.makedirs(dest_folder, exist_ok=True)
            for fname in files:
                if fname.lower().endswith(".pdf"):
                    src = os.path.join(root, fname)
                    dest = os.path.join(dest_folder, fname)
                    try:
                        shutil.copy(src, dest)
                    except PermissionError:
                        print(f"⚠️ Permiso denegado al copiar '{src}'. Se omite.")

        if not ctx.ruta_casos_root:
            ctx.ruta_casos_root = CASOS_DIR_ROOT
            ctx.config_global["casos_path"] = ctx.ruta_casos_root
            guardar_config(ctx.config_global)

        if os.path.isdir(VECTOR_DB_CASOS):
            shutil.rmtree(VECTOR_DB_CASOS)
        os.makedirs(VECTOR_DB_CASOS, exist_ok=True)

        ctx.vectorstores_por_caso.clear()
        ctx.datos_basicos_casos.clear()
        for carpeta in glob.glob(os.path.join(CASOS_DIR_ROOT, "*")):
            if os.path.isdir(carpeta):
                nombre = os.path.basename(carpeta.rstrip(os.sep))
                carpeta_index = os.path.join(VECTOR_DB_CASOS, nombre)
                email_docs = load_email_docs(nombre)
                vs = build_or_load_vectorstore(carpeta, carpeta_index, extra_docs=email_docs)
                ctx.vectorstores_por_caso[nombre] = vs
                if nombre not in ctx.memories_por_caso:
                    ctx.memories_por_caso[nombre] = ChatMessageHistory()
                ctx.datos_basicos_casos[nombre] = extraer_datos_de_carpeta(carpeta)

        carpetas_existentes = [
            os.path.basename(p)
            for p in glob.glob(os.path.join(CASOS_DIR_ROOT, "*"))
            if os.path.isdir(p)
        ]
        return f"✅ Caso '{nombre_carpeta}' cargado e indexado correctamente.", sorted(carpetas_existentes)

    except Exception as e:
        carpetas_existentes = [
            os.path.basename(p)
            for p in glob.glob(os.path.join(CASOS_DIR_ROOT, "*"))
            if os.path.isdir(p)
        ]
        return f"❌ Error al cargar el caso '{nombre_carpeta}': {e}", sorted(carpetas_existentes)


def actualizar_casos(ctx=default_context):
    try:
        if os.path.isdir(VECTOR_DB_CASOS):
            shutil.rmtree(VECTOR_DB_CASOS)
        os.makedirs(VECTOR_DB_CASOS, exist_ok=True)

        ctx.vectorstores_por_caso.clear()
        ctx.datos_basicos_casos.clear()
        carpetas_omitidas = []
        for carpeta in glob.glob(os.path.join(CASOS_DIR_ROOT, "*")):
            if os.path.isdir(carpeta):
                nombre = os.path.basename(carpeta.rstrip(os.sep))
                carpeta_index = os.path.join(VECTOR_DB_CASOS, nombre)
                try:
                    email_docs = load_email_docs(nombre)
                    vs = build_or_load_vectorstore(carpeta, carpeta_index, extra_docs=email_docs)
                except ValueError as e:
                    if "No se encontraron documentos PDF" in str(e):
                        carpetas_omitidas.append(nombre)
                        continue
                    raise
                ctx.vectorstores_por_caso[nombre] = vs
                if nombre not in ctx.memories_por_caso:
                    ctx.memories_por_caso[nombre] = ChatMessageHistory()
                ctx.datos_basicos_casos[nombre] = extraer_datos_de_carpeta(carpeta)
        mensaje = "✅ Índice de Casos actualizado correctamente."
        if carpetas_omitidas:
            carpetas_str = ", ".join(sorted(carpetas_omitidas))
            mensaje += f" ⚠️ Carpetas omitidas sin PDFs: {carpetas_str}"
        return mensaje

    except Exception as e:
        return f"❌ Error al actualizar índice de casos: {e}"


def eliminar_caso(nombre_caso, ctx=default_context):
    """Elimina la carpeta y el índice de un caso existente."""

    nombre = nombre_caso.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de caso a eliminar."

    carpeta_caso = os.path.join(CASOS_DIR_ROOT, nombre)
    carpeta_vector = os.path.join(VECTOR_DB_CASOS, nombre)
    existe = os.path.isdir(carpeta_caso) or os.path.isdir(carpeta_vector)
    try:
        if os.path.isdir(carpeta_caso):
            shutil.rmtree(carpeta_caso)
        if os.path.isdir(carpeta_vector):
            shutil.rmtree(carpeta_vector)
    except Exception as e:
        return f"❌ Error al eliminar el caso '{nombre}': {e}"

    ctx.vectorstores_por_caso.pop(nombre, None)
    ctx.memories_por_caso.pop(nombre, None)
    ctx.datos_basicos_casos.pop(nombre, None)

    if existe:
        return f"✅ Caso '{nombre}' eliminado correctamente."
    else:
        return f"⚠️ El caso '{nombre}' no existe."


def guardar_demandas_por_carpeta(carpeta_sel, ctx=default_context):
    try:
        # Remove previous contents
        for f in glob.glob(os.path.join(DEMANDAS_DIR, "*")):
            try:
                if os.path.isfile(f):
                    os.remove(f)
                else:
                    shutil.rmtree(f)
            except Exception:
                pass

        # Clear previous schemas
        if SCHEMAS_DIR.exists():
            shutil.rmtree(SCHEMAS_DIR)

        # Copy PDFs preserving structure and generate schemas
        for root, _, files in os.walk(carpeta_sel):
            rel = os.path.relpath(root, carpeta_sel)
            dest_folder = os.path.join(DEMANDAS_DIR, rel)
            os.makedirs(dest_folder, exist_ok=True)
            for fname in files:
                if fname.lower().endswith(".pdf"):
                    src = os.path.join(root, fname)
                    dest = os.path.join(dest_folder, fname)
                    try:
                        shutil.copy(src, dest)

                        tipo = _slugify(os.path.splitext(fname)[0])
                        pdf_hash = hash_for_pdf(dest)
                        schema, data = generate_schema_from_pdf(dest)
                        SCHEMAS_DIR.mkdir(parents=True, exist_ok=True)
                        schema_path = SCHEMAS_DIR / f"demanda_{tipo}.json"
                        with open(schema_path, "w", encoding="utf8") as fh:
                            json.dump(schema, fh, indent=2, ensure_ascii=False)
                        cache_form_data(tipo, data)
                        update_schema_index(tipo, pdf_hash)
                    except PermissionError:
                        print(f"⚠️ Permiso denegado al copiar '{src}'. Se omite.")

        ctx.demandas_vectorstore = build_or_load_vectorstore(
            DEMANDAS_DIR, VECTOR_DB_DEMANDAS, force_rebuild=True
        )
        ctx.demandas_textos.clear()
        ctx.demandas_textos.update(cargar_textos_demandas(DEMANDAS_DIR))

        ctx.datos_basicos_demandas = {
            name: extraer_datos_por_llm(txt) or extraer_datos_de_texto(txt)
            for name, txt in ctx.demandas_textos.items()
        }

        ctx.ruta_demandas = carpeta_sel
        ctx.config_global["demandas_path"] = ctx.ruta_demandas
        guardar_config(ctx.config_global)
        return "✅ Demandas de ejemplo cargadas e indexadas correctamente."
    except Exception as e:
        return f"❌ Error al cargar ejemplos de demandas: {e}"


def actualizar_demandas(ctx=default_context):
    src_path = _resolve_path(ctx.ruta_demandas, "")
    if not src_path or not os.path.isdir(src_path):
        return "⚠️ No existe una ruta de ejemplos de demandas válida. Debes cargar primero."
    try:
        ctx.demandas_vectorstore = build_or_load_vectorstore(
            src_path, VECTOR_DB_DEMANDAS, force_rebuild=True
        )
        ctx.demandas_textos.clear()
        ctx.demandas_textos.update(cargar_textos_demandas(src_path))
        ctx.datos_basicos_demandas = {
            name: extraer_datos_por_llm(txt) or extraer_datos_de_texto(txt)
            for name, txt in ctx.demandas_textos.items()
        }
        return "✅ Ejemplos de demandas actualizados correctamente."
    except Exception as e:
        return f"❌ Error al actualizar ejemplos de demandas: {e}"


def eliminar_demandas(ctx=default_context):
    """Elimina las demandas de ejemplo cargadas y su índice vectorial."""
    try:
        if os.path.isdir(DEMANDAS_DIR):
            shutil.rmtree(DEMANDAS_DIR)
        os.makedirs(DEMANDAS_DIR, exist_ok=True)

        if os.path.isdir(VECTOR_DB_DEMANDAS):
            shutil.rmtree(VECTOR_DB_DEMANDAS)
        os.makedirs(VECTOR_DB_DEMANDAS, exist_ok=True)

        ctx.demandas_vectorstore = None
        ctx.demandas_textos.clear()
        ctx.datos_basicos_demandas.clear()
        ctx.ruta_demandas = ""
        if ctx.config_global.get("demandas_path"):
            ctx.config_global["demandas_path"] = ""
            guardar_config(ctx.config_global)
        return "✅ Demandas de ejemplo eliminadas correctamente."
    except Exception as e:
        return f"❌ Error al eliminar demandas: {e}"


def indexar_corpus_legal(ctx=default_context):
    """Construye o actualiza el índice vectorial del corpus legal."""
    try:
        path = _resolve_path(ctx.ruta_legal_corpus, LEGAL_CORPUS_DIR)
        ctx.legal_vectorstore = build_or_load_vectorstore(
            path, VECTOR_DB_LEGAL, force_rebuild=True
        )
        return "✅ Corpus legal indexado correctamente."
    except Exception as e:
        return f"❌ Error al indexar corpus legal: {e}"


def recuperar_textos_legales(pregunta, k=5, ctx=default_context):
    """Recupera documentos relevantes del corpus legal."""
    if not ctx.legal_vectorstore or not pregunta.strip():
        return []
    retriever = ctx.legal_vectorstore.as_retriever(search_kwargs={"k": k})
    try:
        return retriever.get_relevant_documents(pregunta)
    except Exception:
        return []


def _diff_dicts(d1, d2, path=""):
    """Recursively compare two dictionaries and return human readable diffs."""
    diffs = []
    keys = set(d1) | set(d2)
    for k in sorted(keys):
        new_path = f"{path}.{k}" if path else str(k)
        if k not in d1:
            diffs.append(f"+ {new_path}: {d2[k]}")
        elif k not in d2:
            diffs.append(f"- {new_path}: {d1[k]}")
        else:
            v1, v2 = d1[k], d2[k]
            if isinstance(v1, dict) and isinstance(v2, dict):
                diffs.extend(_diff_dicts(v1, v2, new_path))
            elif v1 != v2:
                diffs.append(f"~ {new_path}: {v1} -> {v2}")
    return diffs


def comparar_json_files(path_a, path_b):
    """Return list of differences between two JSON files."""
    with open(path_a, "r", encoding="utf-8") as f:
        data_a = json.load(f)
    with open(path_b, "r", encoding="utf-8") as f:
        data_b = json.load(f)
    return _diff_dicts(data_a, data_b)


# ---------------------------------------------------------------------------
# Nuevas funciones para gestión individual de documentos
# ---------------------------------------------------------------------------

def crear_caso(nombre_caso, ctx=default_context):
    """Crea una carpeta para un caso sin requerir una carpeta de origen."""
    nombre = nombre_caso.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre válido para el caso."

    carpeta_dest = os.path.join(CASOS_DIR_ROOT, nombre)
    if os.path.isdir(carpeta_dest):
        return f"⚠️ El caso '{nombre}' ya existe."

    try:
        os.makedirs(carpeta_dest, exist_ok=True)
        return f"✅ Caso '{nombre}' creado correctamente."
    except Exception as e:
        return f"❌ Error al crear el caso '{nombre}': {e}"


def crear_carpeta_demanda(nombre_carpeta, ctx=default_context):
    """Crea una subcarpeta dentro de ``DEMANDAS_DIR``."""
    nombre = nombre_carpeta.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre válido para la carpeta."

    carpeta_dest = os.path.join(DEMANDAS_DIR, nombre)
    if os.path.isdir(carpeta_dest):
        return f"⚠️ La carpeta '{nombre}' ya existe."

    try:
        os.makedirs(carpeta_dest, exist_ok=True)
        return f"✅ Carpeta de demanda '{nombre}' creada correctamente."
    except Exception as e:
        return f"❌ Error al crear la carpeta '{nombre}': {e}"


def agregar_pdf_a_caso(ruta_pdf, nombre_caso, ctx=default_context):
    """Agrega un PDF individual a un caso existente o nuevo."""
    if not ruta_pdf.lower().endswith(".pdf") or not os.path.isfile(ruta_pdf):
        return "⚠️ Debes seleccionar un archivo PDF válido."

    nombre = nombre_caso.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de caso válido."

    carpeta_dest = os.path.join(CASOS_DIR_ROOT, nombre)
    os.makedirs(carpeta_dest, exist_ok=True)

    try:
        dest = os.path.join(carpeta_dest, os.path.basename(ruta_pdf))
        shutil.copy(ruta_pdf, dest)

        carpeta_index = os.path.join(VECTOR_DB_CASOS, nombre)
        email_docs = load_email_docs(nombre)
        ctx.vectorstores_por_caso[nombre] = build_or_load_vectorstore(
            carpeta_dest, carpeta_index, force_rebuild=True, extra_docs=email_docs
        )
        ctx.datos_basicos_casos[nombre] = extraer_datos_de_carpeta(carpeta_dest)
        if nombre not in ctx.memories_por_caso:
            ctx.memories_por_caso[nombre] = ChatMessageHistory()
        return f"✅ PDF agregado al caso '{nombre}'."
    except Exception as e:
        return f"❌ Error al agregar PDF al caso '{nombre}': {e}"


def guardar_pdf_en_caso(ruta_pdf, nombre_caso):
    """Guarda un adjunto PDF en la carpeta del caso sin indexar."""
    if not ruta_pdf.lower().endswith(".pdf") or not os.path.isfile(ruta_pdf):
        return "⚠️ Debes seleccionar un archivo PDF válido."

    nombre = nombre_caso.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de caso válido."

    carpeta_dest = os.path.join(CASOS_DIR_ROOT, nombre)
    os.makedirs(carpeta_dest, exist_ok=True)

    try:
        dest = os.path.join(carpeta_dest, os.path.basename(ruta_pdf))
        shutil.copy(ruta_pdf, dest)
        return f"✅ PDF guardado en el caso '{nombre}'."
    except Exception as e:
        return f"❌ Error al guardar PDF en el caso '{nombre}': {e}"


def agregar_texto_a_caso(texto, nombre_caso, ctx=default_context):
    """Agrega un texto como archivo .txt a un caso existente o nuevo."""
    if not texto.strip():
        return "⚠️ Debes proporcionar un texto válido."

    nombre = nombre_caso.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de caso válido."

    carpeta_dest = os.path.join(CASOS_DIR_ROOT, nombre)
    os.makedirs(carpeta_dest, exist_ok=True)

    try:
        base = os.path.join(carpeta_dest, "texto")
        contador = 1
        dest = f"{base}.txt"
        while os.path.exists(dest):
            contador += 1
            dest = f"{base}_{contador}.txt"
        with open(dest, "w", encoding="utf-8") as fh:
            fh.write(texto)

        carpeta_index = os.path.join(VECTOR_DB_CASOS, nombre)
        email_docs = load_email_docs(nombre)
        ctx.vectorstores_por_caso[nombre] = build_or_load_vectorstore(
            carpeta_dest,
            carpeta_index,
            force_rebuild=True,
            extra_docs=email_docs,
        )
        ctx.datos_basicos_casos[nombre] = extraer_datos_de_carpeta(carpeta_dest)
        if nombre not in ctx.memories_por_caso:
            ctx.memories_por_caso[nombre] = ChatMessageHistory()
        return f"✅ Texto agregado al caso '{nombre}'."
    except Exception as e:
        return f"❌ Error al agregar texto al caso '{nombre}': {e}"


def guardar_email_en_txt(nombre_caso, nombre, remite, asunto, fecha, cuerpo):
    """Guarda un correo como archivo TXT dentro del caso y devuelve su ruta."""
    nombre_caso = nombre_caso.strip()
    if not nombre_caso:
        raise ValueError("nombre_caso requerido")

    carpeta_dest = os.path.join(CASOS_DIR_ROOT, nombre_caso)
    os.makedirs(carpeta_dest, exist_ok=True)

    slug = _slugify(nombre or asunto or "email") or "email"
    base = os.path.join(carpeta_dest, f"{slug}")
    dest = f"{base}.txt"
    contador = 1
    while os.path.exists(dest):
        contador += 1
        dest = f"{base}_{contador}.txt"

    contenido = (
        f"Remite: {remite}\nAsunto: {asunto}\nFecha: {fecha}\n\n{cuerpo}"
    )
    with open(dest, "w", encoding="utf-8") as fh:
        fh.write(contenido)
    return dest


def agregar_demanda_pdf(ruta_pdf, ctx=default_context, subcarpeta=""):
    """Agrega un PDF a las plantillas de demandas.

    Si ``subcarpeta`` se proporciona, el PDF se copia dentro de dicha
    carpeta para mantener la organización."""
    if not ruta_pdf.lower().endswith(".pdf") or not os.path.isfile(ruta_pdf):
        return "⚠️ Debes seleccionar un archivo PDF válido."

    try:
        dest_folder = os.path.join(DEMANDAS_DIR, subcarpeta)
        os.makedirs(dest_folder, exist_ok=True)
        dest = os.path.join(dest_folder, os.path.basename(ruta_pdf))
        shutil.copy(ruta_pdf, dest)

        # Generate or update schema based on uploaded PDF
        tipo = _slugify(os.path.splitext(os.path.basename(dest))[0])
        pdf_hash = hash_for_pdf(dest)
        schema, data = generate_schema_from_pdf(dest)
        schema_path = SCHEMAS_DIR / f"demanda_{tipo}.json"
        SCHEMAS_DIR.mkdir(parents=True, exist_ok=True)
        with open(schema_path, "w", encoding="utf8") as fh:
            json.dump(schema, fh, indent=2, ensure_ascii=False)
        cache_form_data(tipo, data)
        if index_hash(tipo) != pdf_hash:
            update_schema_index(tipo, pdf_hash)

        ctx.demandas_vectorstore = build_or_load_vectorstore(
            DEMANDAS_DIR, VECTOR_DB_DEMANDAS, force_rebuild=True
        )
        ctx.demandas_textos.clear()
        ctx.demandas_textos.update(cargar_textos_demandas(DEMANDAS_DIR))
        ctx.datos_basicos_demandas = {
            name: extraer_datos_por_llm(txt) or extraer_datos_de_texto(txt)
            for name, txt in ctx.demandas_textos.items()
        }
        return "✅ Demanda agregada e indexada correctamente."
    except Exception as e:
        return f"❌ Error al agregar demanda: {e}"


def eliminar_jurisprudencia_pdf(nombre_pdf, ctx=default_context):
    """Elimina un PDF específico de jurisprudencia."""

    path = os.path.join(JURIS_DIR, nombre_pdf)
    if not os.path.isfile(path):
        return f"⚠️ El archivo '{nombre_pdf}' no existe."
    try:
        os.remove(path)
        restantes = glob.glob(os.path.join(JURIS_DIR, "*.pdf"))
        if restantes:
            ctx.juris_vectorstore = build_or_load_vectorstore(
                JURIS_DIR, VECTOR_DB_JURIS, force_rebuild=True
            )
        else:
            if os.path.isdir(VECTOR_DB_JURIS):
                shutil.rmtree(VECTOR_DB_JURIS)
            ctx.juris_vectorstore = None
        return f"✅ Documento '{nombre_pdf}' eliminado."
    except Exception as e:
        return f"❌ Error al eliminar jurisprudencia: {e}"


def eliminar_demanda_pdf(nombre_pdf, ctx=default_context):
    """Elimina un PDF de las plantillas de demandas.

    ``nombre_pdf`` es la ruta relativa dentro de ``DEMANDAS_DIR``."""
    path = os.path.join(DEMANDAS_DIR, nombre_pdf)
    if not os.path.isfile(path):
        return f"⚠️ El archivo '{nombre_pdf}' no existe."
    try:
        os.remove(path)
        restantes = glob.glob(os.path.join(DEMANDAS_DIR, "**", "*.pdf"), recursive=True)
        if restantes:
            ctx.demandas_vectorstore = build_or_load_vectorstore(
                DEMANDAS_DIR, VECTOR_DB_DEMANDAS, force_rebuild=True
            )
            ctx.demandas_textos.clear()
            ctx.demandas_textos.update(cargar_textos_demandas(DEMANDAS_DIR))
            ctx.datos_basicos_demandas = {
                name: extraer_datos_por_llm(txt) or extraer_datos_de_texto(txt)
                for name, txt in ctx.demandas_textos.items()
            }
        else:
            if os.path.isdir(VECTOR_DB_DEMANDAS):
                shutil.rmtree(VECTOR_DB_DEMANDAS)
            ctx.demandas_vectorstore = None
            ctx.demandas_textos.clear()
            ctx.datos_basicos_demandas.clear()
        return f"✅ Documento '{nombre_pdf}' eliminado."
    except Exception as e:
        return f"❌ Error al eliminar demanda: {e}"


def eliminar_pdf_de_caso(nombre_pdf, nombre_caso, ctx=default_context):
    """Elimina un documento (PDF o TXT) de un caso existente."""
    nombre = nombre_caso.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de caso válido."

    carpeta_caso = os.path.join(CASOS_DIR_ROOT, nombre)
    path = os.path.join(carpeta_caso, nombre_pdf)
    if not os.path.isfile(path):
        return f"⚠️ El archivo '{nombre_pdf}' no existe en el caso '{nombre}'."

    try:
        os.remove(path)
        restantes = glob.glob(os.path.join(carpeta_caso, "**", "*.pdf"), recursive=True)
        restantes += glob.glob(os.path.join(carpeta_caso, "**", "*.txt"), recursive=True)
        carpeta_index = os.path.join(VECTOR_DB_CASOS, nombre)
        if restantes:
            email_docs = load_email_docs(nombre)
            ctx.vectorstores_por_caso[nombre] = build_or_load_vectorstore(
                carpeta_caso,
                carpeta_index,
                force_rebuild=True,
                extra_docs=email_docs,
            )
            ctx.datos_basicos_casos[nombre] = extraer_datos_de_carpeta(carpeta_caso)
        else:
            if os.path.isdir(carpeta_index):
                shutil.rmtree(carpeta_index)
            ctx.vectorstores_por_caso.pop(nombre, None)
            ctx.memories_por_caso.pop(nombre, None)
            ctx.datos_basicos_casos.pop(nombre, None)
        return f"✅ Documento '{nombre_pdf}' eliminado del caso '{nombre}'."
    except Exception as e:
        return f"❌ Error al eliminar PDF del caso '{nombre}': {e}"


def listar_pdfs_de_caso(nombre_caso):
    """Devuelve una lista de documentos (PDF o TXT) de un caso."""
    nombre = nombre_caso.strip()
    if not nombre:
        return []
    carpeta_caso = os.path.join(CASOS_DIR_ROOT, nombre)
    if not os.path.isdir(carpeta_caso):
        return []
    archivos = []
    for patron in ("*.pdf", "*.txt"):
        archivos.extend(glob.glob(os.path.join(carpeta_caso, "**", patron), recursive=True))
    return sorted(os.path.relpath(p, carpeta_caso) for p in archivos)


def crear_area(nombre_area, ctx=default_context):
    """Crea una carpeta para un área."""
    nombre = nombre_area.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre válido para el área."

    carpeta_dest = os.path.join(ctx.ruta_areas_root, nombre)
    if os.path.isdir(carpeta_dest):
        return f"⚠️ El área '{nombre}' ya existe."

    try:
        os.makedirs(carpeta_dest, exist_ok=True)
        return f"✅ Área '{nombre}' creada correctamente."
    except Exception as e:
        return f"❌ Error al crear el área '{nombre}': {e}"


def eliminar_area(nombre_area, ctx=default_context):
    """Elimina la carpeta de un área existente."""
    nombre = nombre_area.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de área a eliminar."

    carpeta_area = os.path.join(ctx.ruta_areas_root, nombre)
    if not os.path.isdir(carpeta_area):
        return f"⚠️ El área '{nombre}' no existe."

    try:
        shutil.rmtree(carpeta_area)
        return f"✅ Área '{nombre}' eliminada correctamente."
    except Exception as e:
        return f"❌ Error al eliminar el área '{nombre}': {e}"


def agregar_pdf_a_area(ruta_pdf, nombre_area, ctx=default_context):
    """Agrega un PDF a una carpeta de área."""
    if not ruta_pdf.lower().endswith(".pdf") or not os.path.isfile(ruta_pdf):
        return "⚠️ Debes seleccionar un archivo PDF válido."

    nombre = nombre_area.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de área válido."

    carpeta_dest = os.path.join(ctx.ruta_areas_root, nombre)
    os.makedirs(carpeta_dest, exist_ok=True)

    try:
        dest = os.path.join(carpeta_dest, os.path.basename(ruta_pdf))
        shutil.copy(ruta_pdf, dest)
        return f"✅ PDF agregado al área '{nombre}'."
    except Exception as e:
        return f"❌ Error al agregar PDF al área '{nombre}': {e}"


def eliminar_pdf_de_area(nombre_pdf, nombre_area, ctx=default_context):
    """Elimina un PDF de un área existente."""
    nombre = nombre_area.strip()
    if not nombre:
        return "⚠️ Debes indicar un nombre de área válido."

    carpeta_area = os.path.join(ctx.ruta_areas_root, nombre)
    path = os.path.join(carpeta_area, nombre_pdf)
    if not os.path.isfile(path):
        return f"⚠️ El archivo '{nombre_pdf}' no existe en el área '{nombre}'."

    try:
        os.remove(path)
        return f"✅ Documento '{nombre_pdf}' eliminado del área '{nombre}'."
    except Exception as e:
        return f"❌ Error al eliminar PDF del área '{nombre}': {e}"


def mover_pdf_de_caso_a_area(nombre_pdf, nombre_caso, nombre_area, ctx=default_context):
    """Mueve un PDF de un caso a un área."""
    ruta = os.path.join(CASOS_DIR_ROOT, nombre_caso, nombre_pdf)
    if not os.path.isfile(ruta):
        return f"⚠️ El archivo '{nombre_pdf}' no existe en el caso '{nombre_caso}'."

    msg = agregar_pdf_a_area(ruta, nombre_area, ctx)
    if not msg.startswith("✅"):
        return msg

    msg_del = eliminar_pdf_de_caso(nombre_pdf, nombre_caso, ctx)
    if msg_del.startswith("✅"):
        return f"✅ Documento '{nombre_pdf}' movido al área '{nombre_area}'."
    return msg_del


def listar_pdfs_de_area(nombre_area, ctx=default_context):
    """Devuelve una lista de documentos (PDF o TXT) de un área."""
    nombre = nombre_area.strip()
    if not nombre:
        return []
    carpeta_area = os.path.join(ctx.ruta_areas_root, nombre)
    if not os.path.isdir(carpeta_area):
        return []
    archivos = []
    for patron in ("*.pdf", "*.txt"):
        archivos.extend(
            glob.glob(os.path.join(carpeta_area, "**", patron), recursive=True)
        )
    return sorted(os.path.relpath(p, carpeta_area) for p in archivos)


def listar_areas(ctx=default_context):
    """Return a list of available area names, including nested folders."""
    carpetas = []
    for root, dirs, _ in os.walk(ctx.ruta_areas_root):
        for d in dirs:
            path = os.path.join(root, d)
            carpetas.append(os.path.relpath(path, ctx.ruta_areas_root))
    return sorted(carpetas)


# ---------------------------------------------------------------------------
# Utilidades para el formulario guiado de demandas
# ---------------------------------------------------------------------------

def sugerir_juzgados(tipo_caso="general"):
    """Devuelve una lista básica de juzgados sugeridos según el tipo."""
    sugerencias = {
        "laboral": ["Unidad Judicial Laboral", "Tribunal Distrital de lo Laboral"],
        "civil": ["Unidad Judicial Civil", "Tribunal Provincial"],
        "familia": ["Unidad Judicial de Familia"],
        "general": ["Unidad Judicial Multicompetente"],
    }
    return sugerencias.get(tipo_caso, sugerencias["general"])


def sugerir_fundamentos_derecho(tipo_accion="general"):
    """Devuelve una lista básica de artículos sugeridos según el tipo."""
    sugerencias = {
        "laboral": ["Art. 42 Código del Trabajo", "Art. 575 COGEP"],
        "civil": ["Art. 142 COGEP", "Art. 27 Constitución"],
        "familia": ["Art. 67 Constitución"],
        "general": ["Art. 27 Constitución", "Art. 130 COGEP"],
    }
    return sugerencias.get(tipo_accion, sugerencias["general"])


def elaborar_fundamentos_derecho(caso=None, ctx=default_context):
    """Elabora el numeral 6 de la demanda citando normas y precedentes."""
    pregunta = (
        "Elabora el numeral 6 (Fundamentos de Derecho) de una demanda, "
        "citando con exactitud las normas y precedentes que sustenten la acción."
    )
    if caso:
        pregunta += f" Caso: {caso}."
    texto = buscar_palabras_clave_fn(pregunta, ctx)
    if texto:
        try:
            texto = mejorar_fundamentos_llm(texto, ctx)
        except Exception:
            pass
    return texto


def determinar_procedimiento(cuantia, materia):
    """Determina el procedimiento a seguir basado en la cuantía."""
    try:
        valor = float(cuantia)
    except Exception:
        return ""
    if valor <= 10000:
        return "Sumario"
    return "Ordinario"


def generar_formula_pretension():
    """Devuelve una fórmula estándar de pretensión."""
    return (
        "En virtud de los hechos y fundamentos de derecho expuestos, "
        "solicito se sirva admitir la presente demanda y se conceda todo lo "
        "peticionado conforme a derecho."
    )


def sugerir_datos_para_formulario(caso_seleccionado, ctx=default_context):
    """Obtiene un diccionario de valores sugeridos para el formulario."""
    if not caso_seleccionado:
        return {}

    datos = {}
    for sec in SECCIONES_DEMANDA:
        if sec == "DATOS_ACTOR":
            for campo in CAMPOS_DATOS_ACTOR:
                val = obtener_dato_de_caso(campo, caso_seleccionado, ctx)
                if val:
                    datos[campo] = val
            continue

        if sec == "DATOS_DEFENSOR":
            for campo in CAMPOS_DATOS_DEFENSOR:
                val = obtener_dato_de_caso(campo, caso_seleccionado, ctx)
                if val:
                    datos[campo] = val
            continue

        if sec == "DATOS_DEMANDADO":
            for campo in CAMPOS_DATOS_DEMANDADO:
                val = obtener_dato_de_caso(campo, caso_seleccionado, ctx)
                if val:
                    datos[campo] = val
            continue

        val = obtener_dato_de_caso(sec, caso_seleccionado, ctx)
        if val:
            datos[sec] = val

    if "HECHOS" not in datos:
        resumen = resumir_caso(caso_seleccionado, ctx)
        if resumen:
            datos["HECHOS"] = resumen

    if ctx.juris_vectorstore:
        for sec in ("FUNDAMENTOS_DERECHO", "PRETENSION"):
            if sec not in datos:
                try:
                    pregunta = (
                        f"¿Cuáles son los {sec.replace('_', ' ').lower()} para el caso {caso_seleccionado}?"
                    )
                    sugerido = buscar_palabras_clave_fn(pregunta, ctx)
                    if sugerido:
                        datos[sec] = sugerido
                except Exception:
                    pass
    return datos


def generar_resumen_demanda(caso_seleccionado, ctx=default_context):
    """Devuelve un resumen ordenado con los datos del formulario de demanda."""
    datos = sugerir_datos_para_formulario(caso_seleccionado, ctx)
    secciones = []

    def _vacio(val):
        return val if val else "VACIO"

    secciones.append("PRIMERO. - DESIGNACIÓN DEL JUZGADOR")
    secciones.append(_vacio(datos.get("DESIGNACION_JUZGADOR")))
    secciones.append("")

    secciones.append("SEGUNDO. - DATOS DEL ACTOR")
    for campo in CAMPOS_DATOS_ACTOR:
        etiqueta = campo.replace("_", " ").title()
        secciones.append(f"{etiqueta}: {_vacio(datos.get(campo))}")
    secciones.append("")

    secciones.append("TERCERO. - DATOS DEL DEFENSOR")
    for campo in CAMPOS_DATOS_DEFENSOR:
        etiqueta = campo.replace("_", " ").title()
        secciones.append(f"{etiqueta}: {_vacio(datos.get(campo))}")
    secciones.append("")

    secciones.append("CUARTO. - RUC")
    secciones.append(_vacio(datos.get("RUC")))
    secciones.append("")

    secciones.append("QUINTO. - DATOS DEL DEMANDADO")
    for campo in CAMPOS_DATOS_DEMANDADO:
        etiqueta = campo.replace("_", " ").title()
        secciones.append(f"{etiqueta}: {_vacio(datos.get(campo))}")
    secciones.append("")

    secciones.append("SEXTO. - HECHOS")
    secciones.append(_vacio(datos.get("HECHOS")))
    secciones.append("")

    secciones.append("SÉPTIMO. - FUNDAMENTOS DE DERECHO")
    secciones.append(_vacio(datos.get("FUNDAMENTOS_DERECHO")))
    secciones.append("")

    secciones.append("OCTAVO. - ACCESO A PRUEBAS")
    secciones.append(_vacio(datos.get("ACCESO_PRUEBAS")))
    secciones.append("")

    secciones.append("NOVENO. - PRETENSIÓN")
    secciones.append(_vacio(datos.get("PRETENSION")))
    secciones.append("")

    secciones.append("DÉCIMO. - CUANTÍA")
    secciones.append(_vacio(datos.get("CUANTIA")))
    secciones.append("")

    secciones.append("DÉCIMO PRIMERO. - PROCEDIMIENTO")
    secciones.append(_vacio(datos.get("PROCEDIMIENTO")))
    secciones.append("")

    secciones.append("DÉCIMO SEGUNDO. - FIRMAS")
    secciones.append(_vacio(datos.get("FIRMAS")))

    return "\n".join(secciones).strip()


def generar_redaccion_demanda(datos):
    """Devuelve un texto redactado en tono jurídico con los datos del formulario."""
    def _vacio(val):
        return val if val else "VACIO"

    partes = []
    partes.append("SEÑOR/A JUEZ/A:")

    seccion = 1

    partes.append(f"{_ordinal(seccion)}. - LA DESIGNACIÓN DEL JUEZ ANTE QUIEN SE PROPONE.")
    partes.append(_vacio(datos.get("DESIGNACION_JUZGADOR")))
    partes.append("El juez ante quien se propone esta demanda queda designado.")
    partes.append("")
    seccion += 1

    partes.append(
        f"{_ordinal(seccion)}. - LOS NOMBRES Y APELLIDOS Y MÁS GENERALES DE LEY DEL ACTOR SON:"
    )
    actor = datos.get("ACTOR_NOMBRES_APELLIDOS")
    if actor:
        linea = f"{actor}"
        ced = datos.get("ACTOR_CEDULA") or datos.get("ACTOR_PASAPORTE")
        if ced:
            linea += f", portador de la cédula {ced}"
        estado = datos.get("ACTOR_ESTADO_CIVIL")
        if estado:
            linea += f", {estado}"
        edad = datos.get("ACTOR_EDAD")
        if edad:
            linea += f", de {edad} años"
        profesion = datos.get("ACTOR_PROFESION")
        if profesion:
            linea += f", de profesión {profesion}"
        dom_parts = [
            datos.get("ACTOR_PROVINCIA"),
            datos.get("ACTOR_CANTON"),
            datos.get("ACTOR_CALLE_PRIMARIA"),
            datos.get("ACTOR_CALLE_SECUNDARIA"),
            datos.get("ACTOR_NUMERO_CASA"),
        ]
        dom = ", ".join(p for p in dom_parts if p)
        if dom:
            linea += f", domiciliado en {dom}"
        email = datos.get("ACTOR_DIR_ELECTRONICA")
        if email:
            linea += f", correo {email}"
        partes.append(linea)
    else:
        partes.append("VACIO")
    partes.append(
        "ante usted comparezco con la siguiente DEMANDA EJECUTIVA."
    )
    partes.append("")
    seccion += 1

    ruc = datos.get("RUC")
    partes.append(f"{_ordinal(seccion)}. - Mi número de RUC es {_vacio(ruc)}.")
    partes.append("")
    seccion += 1

    demandado = datos.get("DEMANDADO_NOMBRES_APELLIDOS")
    if demandado:
        linea = f"Demando al señor {demandado}"
        dced = datos.get("DEMANDADO_CEDULA")
        if dced:
            linea += f", con cédula {dced}"
        dnac = datos.get("DEMANDADO_NACIONALIDAD")
        if dnac:
            linea += f", de nacionalidad {dnac}"
        dprof = datos.get("DEMANDADO_PROFESION")
        if dprof:
            linea += f", profesión {dprof}"
        dedad = datos.get("DEMANDADO_EDAD")
        if dedad:
            linea += f", {dedad} años de edad"
        ddom_parts = [
            datos.get("DEMANDADO_PROVINCIA"),
            datos.get("DEMANDADO_CANTON"),
            datos.get("DEMANDADO_CALLE_PRIMARIA"),
            datos.get("DEMANDADO_CALLE_SECUNDARIA"),
            datos.get("DEMANDADO_NUMERO_CASA"),
        ]
        ddom = ", ".join(p for p in ddom_parts if p)
        if ddom:
            linea += f", domiciliado en {ddom}"
        desc = datos.get("DEMANDADO_DESCRIPCION_VIVIENDA")
        if desc:
            linea += f", {desc}"
        dmail = datos.get("DEMANDADO_DIR_ELECTRONICA")
        if dmail:
            linea += f", correo electrónico {dmail}"
        linea += "."
        partes.append(f"{_ordinal(seccion)}. - {linea}")
    else:
        partes.append(f"{_ordinal(seccion)}. - VACIO")
    partes.append("")
    seccion += 1

    defensor = datos.get("DEFENSOR_NOMBRE")
    if defensor:
        casillero = datos.get("CASILLERO_JUDICIAL", "")
        partes.append(
            f"{_ordinal(seccion)}. - Actúo patrocinado por el abogado {defensor}, casillero judicial {casillero}."
        )
    else:
        partes.append(f"{_ordinal(seccion)}. - VACIO")
    partes.append("")
    seccion += 1

    partes.append(f"{_ordinal(seccion)}. - HECHOS:")
    partes.append(_vacio(datos.get("HECHOS")))
    partes.append("")
    seccion += 1

    partes.append(f"{_ordinal(seccion)}. - FUNDAMENTOS DE DERECHO:")
    partes.append(_vacio(datos.get("FUNDAMENTOS_DERECHO")))
    partes.append("")
    seccion += 1

    partes.append(f"{_ordinal(seccion)}. - ACCESO A PRUEBAS:")
    partes.append(_vacio(datos.get("ACCESO_PRUEBAS")))
    partes.append("")
    seccion += 1

    partes.append(f"{_ordinal(seccion)}. - PRETENSIÓN:")
    partes.append(_vacio(datos.get("PRETENSION")))
    partes.append("")
    seccion += 1

    partes.append(
        f"{_ordinal(seccion)}. - La cuantía del proceso asciende a {_vacio(datos.get('CUANTIA'))}."
    )
    partes.append("")
    seccion += 1

    partes.append(
        f"{_ordinal(seccion)}. - El procedimiento aplicable es el {_vacio(datos.get('PROCEDIMIENTO'))}."
    )
    partes.append("")
    seccion += 1

    partes.append(f"{_ordinal(seccion)}. - " + _vacio(datos.get("FIRMAS")))
    return "\n".join(partes).strip()



def generar_redaccion_demanda_llm(datos, llm_model=None, caso=None, ctx=default_context):
    """Genera una demanda estructurada en 12 puntos.

    Se utilizan los datos del formulario y, si algún valor falta, se intenta
    completar con sugerencias del ``caso`` indicado. Las sugerencias se
    muestran con la etiqueta ``(Sugerido)`` para indicar que provienen del LLM."""
    sugeridos = sugerir_datos_para_formulario(caso, ctx) if caso else {}

    def _valor(clave):
        val = datos.get(clave, "").strip()
        if val:
            return val
        val = sugeridos.get(clave, "")
        if val:
            return f"{val} (Sugerido)"
        return "(Sugerido)"

    completos = {}
    for campo in CAMPOS_DATOS_ACTOR + CAMPOS_DATOS_DEFENSOR + CAMPOS_DATOS_DEMANDADO:
        completos[campo] = _valor(campo)

    for sec in (
        "DESIGNACION_JUZGADOR",
        "RUC",
        "HECHOS",
        "FUNDAMENTOS_DERECHO",
        "ACCESO_PRUEBAS",
        "PRETENSION",
        "CUANTIA",
        "PROCEDIMIENTO",
        "FIRMAS",
    ):
        completos[sec] = _valor(sec)

    if completos.get("HECHOS"):
        try:
            completos["HECHOS"] = mejorar_hechos_llm(completos["HECHOS"], llm_model)
        except Exception:
            pass

    if completos.get("FUNDAMENTOS_DERECHO"):
        try:
            completos["FUNDAMENTOS_DERECHO"] = mejorar_fundamentos_llm(
                completos["FUNDAMENTOS_DERECHO"], ctx, llm_model
            )
        except Exception:
            pass
    refs = []
    consulta = "\n".join(
        [completos.get("HECHOS", ""), completos.get("FUNDAMENTOS_DERECHO", "")]
    )
    if consulta.strip():
        refs = recuperar_textos_legales(consulta, ctx=ctx)

    texto = generar_redaccion_demanda(completos)
    if refs:
        citas = "\n\nReferencias:\n" + "\n".join(
            f"- {d.metadata.get('source', 'desconocido')}" for d in refs
        )
        texto += citas
    return texto


def armar_reemplazos(datos, plantilla):
    """Devuelve un diccionario de reemplazos para ``plantilla``.

    - Las líneas de ``datos['HECHOS']`` se enumeran como ``HECHO1``, ``HECHO2``,
      etc.
    - Se agregan claves con guiones bajos para cualquier marcador presente en la
      plantilla que no tenga un valor correspondiente.
    """
    reemplazos = {k: v for k, v in datos.items() if k != "HECHOS"}

    hechos = datos.get("HECHOS", "")
    lineas = [ln.strip() for ln in hechos.splitlines() if ln.strip()]
    for i, linea in enumerate(lineas, 1):
        reemplazos[f"HECHO{i}"] = linea

    marcadores = set(re.findall(r"\[([A-Z0-9_]+)\]", plantilla))
    for ph in marcadores:
        reemplazos.setdefault(ph, "___")

    return reemplazos



def mejorar_hechos_llm(texto, llm_model=None):
    """Devuelve cada línea de ``texto`` mejorada con ayuda de un LLM."""
    lineas = [ln.strip() for ln in texto.splitlines() if ln.strip()]
    if not lineas:
        return texto
    llm_model = llm_model or get_llm()
    prompt = ChatPromptTemplate.from_messages([
        (
            "system",
            "Mejora la redacción del siguiente hecho de una demanda de forma breve y clara."
        ),
        ("human", "{hecho}")
    ])
    chain = prompt | llm_model | StrOutputParser()
    nuevas = []
    for linea in lineas:
        try:
            resp = chain.invoke({"hecho": linea}).strip()
            resp = "\n".join(
                ln
                for ln in resp.splitlines()
                if not re.match(r"\s*(comentario|comment):", ln, re.IGNORECASE)
            ).strip()
            nuevas.append(resp if resp else linea)
        except Exception:
            nuevas.append(linea)
    return "\n".join(nuevas)


def mejorar_fundamentos_llm(texto, ctx=default_context, llm_model=None):
    """Verifica y mejora los fundamentos citando correctamente los artículos."""
    if not texto.strip() or not ctx.juris_vectorstore:
        return texto
    llm_model = llm_model or get_llm()
    retr = ctx.juris_vectorstore.as_retriever(search_kwargs={"k": 25})
    raw = ConversationalRetrievalChain.from_llm(
        llm=llm_model,
        retriever=retr,
        combine_docs_chain_kwargs={"prompt": prompt},
    )
    chain = RunnableWithMessageHistory(
        raw,
        lambda _sid: ctx.memory_palabras,
        input_messages_key="question",
        history_messages_key="chat_history",
    )
    pregunta = (
        "Revisa y corrige, citando con precisión los artículos legales, el siguiente texto:"\
        f"\n{texto}"
    )
    try:
        resp = chain.invoke({"question": pregunta}, config={"configurable": {"session_id": "fundamentos"}})["answer"].strip()
        return resp if resp else texto
    except Exception:
        return texto
